"""
Grandiose Bakery — Financial Performance & Cost Optimization Dashboard
GIP III | Parineeta Jain, Rajveer Singh, Tarang Gupta

Bold dark-theme redesign: near-black canvas, vibrant amber/coral/violet
accents, sparkline KPI cards, glowing trend charts, pill navigation.

Two tabs:
  Tab A - Performance tracker (KPIs, trend, category panels)
  Tab B - Scenario & resilience (4 modules)

Scope: bakery operations only (catering excluded per GIP III scope).
"""

import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import smtplib
import re
from datetime import datetime
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from io import BytesIO
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
import json
import math
from anthropic import Anthropic
from pypdf import PdfReader
from docx import Document as DocxDocument
from supabase import create_client
import sku_data
import labour_calc
from labour_config import LABOUR_CONFIG

# ----------------------------------------------------------------------
# PAGE CONFIG
# ----------------------------------------------------------------------
st.set_page_config(
    page_title="Grandiose Bakery Dashboard",
    page_icon="🍞",
    layout="wide",
    # Collapsed by default: navigation lives in the top bar, so the canvas
    # runs full-width like the design. The sidebar is still there for the
    # email-report tool and project credits.
    initial_sidebar_state="collapsed",
)

# ----------------------------------------------------------------------
# COLOR PALETTE — bold dark theme. Near-black canvas with vibrant,
# bakery-warm accents (amber/gold as primary, coral and violet as
# secondary/tertiary) instead of generic blue/purple.
# ----------------------------------------------------------------------
COLORS = {
    # Exact tokens from Stitch's "Grandiose Cinematic Dashboard" DESIGN.md
    "bg":        "#0A0A0B",   # Level 0 background
    "surface_lowest": "#0E0E0F",  # below-background band — footer, app chrome
    "surface":   "#161618",   # Level 1 card surface
    "surface2":  "#1C1B1C",   # sidebar / secondary surface (surface-container-low)
    "border":    "#2A2A2C",   # card border
    "text":      "#E5E2E3",   # on-surface, warm cream — headings
    "text_soft": "#D0C5AF",   # on-surface-variant, muted warm tan — secondary text/labels
    "primary":   "#F2CA50",   # primary gold — interactive states, positive metrics
    "primary_deep": "#D4AF37",  # deeper gold — chart lines/gradients/dividers
    "primary_d": "#3C2F00",   # on-primary — dark text for solid-gold CTA buttons
    "secondary": "#C6C7C2",   # light gray, used sparingly
    "tertiary":  "#8A8370",   # muted tan-gray, for chart variety
    "success":   "#F2CA50",   # gold = positive, per this system's own semantics
    "warning":   "#B8AE8E",   # muted khaki — caution status
    "danger":    "#FFB4AB",   # error/risk status, straight from the design system's error token
}

CATEGORICAL = [COLORS["primary"], COLORS["primary_deep"], COLORS["text"], COLORS["text_soft"], COLORS["secondary"]]

def hex_to_rgba(hex_color, alpha):
    hex_color = hex_color.lstrip("#")
    r, g, b = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    return f"rgba({r},{g},{b},{alpha})"


# Web fonts MUST be pulled in with an @import inside a <style> block, not a
# <link> tag. st.markdown injects HTML via innerHTML, and browsers treat a
# <link> inserted that way as inert — it lands in the DOM but is never
# fetched, so every heading silently fell back to a system font. A <style>
# element inserted the same way *is* parsed, and its @import does get
# fetched. @import has to be the first rule in its own stylesheet, which is
# why this sits in a block of its own.
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Bebas+Neue&family=Inter:wght@400;500;600&display=swap');
</style>
""", unsafe_allow_html=True)

st.markdown(f"""
<style>
    .stApp {{ background-color: {COLORS['bg']}; }}
    section[data-testid="stSidebar"] {{ background-color: {COLORS['surface2']}; border-right: 1px solid {COLORS['border']}; }}

    /* Top app header/toolbar band (Share, star, edit, GitHub icons) —
       white by default, doesn't inherit the dark theme on its own.
       Real testid is "stHeader" (className happens to be "stAppHeader",
       which is not the same thing — that mismatch was the bug last time). */
    [data-testid="stHeader"] {{
        background-color: {COLORS['bg']} !important;
        border-bottom: 1px solid {COLORS['border']} !important;
    }}
    [data-testid="stHeader"] * {{
        color: {COLORS['text']} !important;
        fill: {COLORS['text']} !important;
    }}
    [data-testid="stHeader"] [data-testid^="stBaseButton"] {{
        background-color: {COLORS['surface']} !important;
    }}
    [data-testid="stHeader"] [data-testid^="stBaseButton"] p,
    [data-testid="stHeader"] [data-testid^="stBaseButton"] span {{
        color: {COLORS['text']} !important;
    }}
    [data-testid="stMainMenu"], [data-testid="stStatusWidget"], [data-testid="stAppDeployButton"] {{
        background-color: transparent !important;
        color: {COLORS['text']} !important;
        fill: {COLORS['text']} !important;
    }}
    /* The actual clickable icon is a nested "stMainMenuButton", a different
       component from the other header buttons — this is the specific
       element that was still showing as an unstyled white block. */
    [data-testid="stMainMenuButton"] {{
        background-color: transparent !important;
        color: {COLORS['text']} !important;
    }}
    [data-testid="stMainMenuButton"] svg, [data-testid="stMainMenuButton"] path,
    [data-testid="stMainMenu"] svg, [data-testid="stStatusWidget"] svg {{
        fill: {COLORS['text']} !important;
        color: {COLORS['text']} !important;
    }}

    /* File uploader — large centered dashed drop zone, button-forward like
       a "Select files / or drop files here" upload box, in our dark theme. */
    [data-testid="stFileUploaderDropzone"] {{
        background-color: {COLORS['surface']} !important;
        border: 2px dashed {hex_to_rgba(COLORS['primary'], 0.55)} !important;
        border-radius: 4px !important;
        padding: 44px 24px !important;
        flex-direction: column-reverse !important;
        align-items: center !important;
        justify-content: center !important;
        text-align: center !important;
        gap: 14px !important;
        min-height: 200px !important;
        transition: border-color 0.15s ease, background-color 0.15s ease;
    }}
    [data-testid="stFileUploaderDropzone"]:hover {{
        border-color: {COLORS['primary']} !important;
        background-color: {COLORS['surface2']} !important;
    }}
    [data-testid="stFileUploaderDropzoneInstructions"],
    [data-testid="stFileUploaderDropzoneInstructions"] * {{
        color: {COLORS['text_soft']} !important;
        justify-content: center !important;
    }}
    [data-testid="stFileUploaderDropzone"] svg {{
        fill: {COLORS['text_soft']} !important;
        width: 32px !important; height: 32px !important;
    }}
    [data-testid="stFileUploaderDropzone"] [data-testid^="stBaseButton"] {{
        padding: 12px 32px !important;
        font-size: 1rem !important;
        border-radius: 4px !important;
    }}
    [data-testid="stFileUploaderFile"] {{
        background-color: {COLORS['surface2']} !important;
        border-radius: 4px !important;
        border: 1px solid {COLORS['border']} !important;
    }}
    [data-testid="stFileUploaderFile"] * {{
        color: {COLORS['text']} !important;
    }}

    .stApp, .stApp p, .stApp span, .stApp li, .stApp label,
    div[data-testid="stMarkdownContainer"], div[data-testid="stMarkdownContainer"] p {{
        color: {COLORS['text']} !important;
        font-family: 'Inter', sans-serif !important;
    }}
    div[data-testid="stCaptionContainer"], .stApp small {{ color: {COLORS['text_soft']} !important; }}
    /* Headings are Bebas Neue — but Streamlit renders every heading as
       <h1><span>text</span>…</h1>, and the broad `.stApp span` Inter rule
       above matches that inner span directly, which outranks a bare `h1`
       rule and silently forced every heading in the app back to Inter.
       Targeting the span too (and scoping with .stApp so these selectors
       outweigh `.stApp span`) is what actually makes Bebas render. */
    .stApp h1, .stApp h2,
    .stApp h1 span, .stApp h2 span {{
        color: {COLORS['text']} !important; font-family: 'Bebas Neue', sans-serif !important;
        font-weight: 400 !important; letter-spacing: -0.01em; text-transform: uppercase;
    }}
    .stApp h3, .stApp h4, .stApp h5,
    .stApp h3 span, .stApp h4 span, .stApp h5 span {{
        color: {COLORS['text']} !important; font-family: 'Bebas Neue', sans-serif !important;
        font-weight: 400 !important; letter-spacing: 0.04em; text-transform: uppercase;
    }}
    /* The heading anchor-link affordance holds an SVG, not text — keep it
       out of the heading type treatment. */
    .stApp [data-testid="stHeaderActionElements"] {{ text-transform: none; }}
    section[data-testid="stSidebar"] * {{ color: {COLORS['text']} !important; font-family: 'Inter', sans-serif !important; }}

    /* Restore Streamlit's icon-ligature font wherever our broad font-family
       rules above would otherwise break it — this is what renders expander/
       popover arrows correctly instead of as literal "expand_more" text. */
    span[data-testid="stIconMaterial"], [data-testid="stIconMaterial"] * {{
        font-family: 'Material Symbols Rounded' !important;
    }}
    section[data-testid="stSidebar"] span[data-testid="stIconMaterial"] {{
        font-family: 'Material Symbols Rounded' !important;
    }}

    /* Popovers render in a separate layer that doesn't inherit .stApp's dark
       background — without this, forced near-white text sits on Streamlit's
       default white popover background and becomes unreadable. */
    div[data-testid="stPopoverBody"] {{
        background-color: {COLORS['surface']} !important;
        border: 1px solid {COLORS['border']} !important;
        border-radius: 4px !important;
    }}
    div[data-testid="stPopoverBody"] * {{
        color: {COLORS['text']} !important;
    }}
    div[data-testid="stPopoverBody"] input, div[data-testid="stPopoverBody"] textarea {{
        background-color: {COLORS['surface2']} !important;
        color: {COLORS['text']} !important;
        border: 1px solid {COLORS['border']} !important;
    }}
    div[data-testid="stPopoverBody"] ::placeholder {{
        color: {COLORS['text_soft']} !important;
        opacity: 1 !important;
    }}

    /* Every floating/portaled element (tooltips, selectbox dropdown menus,
       date pickers, multiselect lists) gets rendered by baseweb OUTSIDE
       .stApp, in a "baseweb-layer" portal attached to <body> — none of it
       inherits our dark theme by default, which is why hover tooltips and
       dropdown option lists have shown up as white-on-white. Fix them all
       at once rather than chasing each one individually. */
    .baseweb-layer, [data-testid="stTooltipContent"], [data-baseweb="menu"],
    [data-baseweb="popover"], [data-baseweb="calendar"] {{
        background-color: {COLORS['surface']} !important;
        border: 1px solid {COLORS['border']} !important;
    }}
    .baseweb-layer *, [data-testid="stTooltipContent"] *, [data-baseweb="menu"] *,
    [data-baseweb="popover"] *, [data-baseweb="calendar"] * {{
        color: {COLORS['text']} !important;
        background-color: transparent !important;
    }}
    [data-baseweb="menu"] li:hover, [data-baseweb="menu"] li[aria-selected="true"] {{
        background-color: {hex_to_rgba(COLORS['primary'], 0.18)} !important;
    }}

    /* Buttons, per the DESIGN.md: "Primary buttons are ghost-style with a
       Gold border and Gold label-sm text. Fill buttons are only used for
       the most critical Call to Action." So every button defaults to ghost,
       and only Streamlit's type="primary" buttons get the solid gold fill —
       reserved deliberately for the handful of true CTAs (Submit, Send). */
    [data-testid^="stBaseButton"] {{
        background-color: transparent !important;
        border: 1px solid {COLORS['text_soft']} !important;
        border-radius: 4px !important;
        transition: border-color 0.2s ease, background-color 0.2s ease;
    }}
    [data-testid^="stBaseButton"] p, [data-testid^="stBaseButton"] span,
    [data-testid^="stBaseButton"] div {{
        color: {COLORS['primary']} !important;
        font-weight: 600 !important;
        text-transform: uppercase; font-size: 0.78rem; letter-spacing: 0.08em;
    }}
    [data-testid^="stBaseButton"]:hover {{
        border-color: {COLORS['primary']} !important;
        background-color: {hex_to_rgba(COLORS['primary_deep'], 0.08)} !important;
    }}
    [data-testid^="stBaseButton"]:disabled, [data-testid^="stBaseButton"]:disabled p {{
        background-color: transparent !important;
        border-color: {COLORS['border']} !important;
        color: {COLORS['text_soft']} !important;
    }}
    /* The critical-CTA exception: solid gold fill, dark text, no border. */
    [data-testid*="rimary" i] {{
        background-color: {COLORS['primary']} !important;
        border: none !important;
    }}
    [data-testid*="rimary" i] p, [data-testid*="rimary" i] span, [data-testid*="rimary" i] div {{
        color: {COLORS['primary_d']} !important;
    }}
    [data-testid*="rimary" i]:hover {{
        background-color: {COLORS['primary_deep']} !important;
    }}

    /* Expander headers — safety net so the label never fades to low-contrast */
    div[data-testid="stExpander"] summary p, div[data-testid="stExpander"] summary span {{
        color: {COLORS['text']} !important;
        opacity: 1 !important;
    }}


    /* Bordered containers -> "glow-card" treatment from the DESIGN.md:
       tiered surface, a top-edge gold hairline, backdrop blur. Per spec,
       hovering does NOT lift the card — only the border and glow intensify.

       Streamlit renamed this container's test-id, so the original selector
       matched nothing and every card in the app quietly fell back to the
       plain default border with no surface tint, hairline or blur. Both the
       old and current structures are listed so this survives a version
       change either way. The :not([class*="st-key-"]) guard keeps
       purely-structural keyed containers (the top bar, the panel row) from
       being painted as cards — only real st.container(border=True) cards
       are, since those carry no key. */
    div[data-testid="stVerticalBlockBorderWrapper"],
    [data-testid="stLayoutWrapper"] > [data-testid="stVerticalBlock"]:not([class*="st-key-"]) {{
        background-color: {hex_to_rgba(COLORS['surface'], 0.85)};
        backdrop-filter: blur(12px);
        border: 1px solid {COLORS['border']} !important;
        border-radius: 4px !important;
        position: relative;
        transition: border-color 0.3s ease, box-shadow 0.3s ease;
    }}
    div[data-testid="stVerticalBlockBorderWrapper"]::before,
    [data-testid="stLayoutWrapper"] > [data-testid="stVerticalBlock"]:not([class*="st-key-"])::before {{
        content: '';
        position: absolute; top:0; left:0; right:0; height:1px;
        background: linear-gradient(90deg, transparent, {hex_to_rgba(COLORS['primary_deep'], 0.15)}, transparent);
        pointer-events: none;
        border-radius: 4px 4px 0 0;
    }}
    div[data-testid="stVerticalBlockBorderWrapper"]:hover,
    [data-testid="stLayoutWrapper"] > [data-testid="stVerticalBlock"]:not([class*="st-key-"]):hover {{
        border-color: #3F3F42 !important;
        box-shadow: 0 10px 30px -10px rgba(0,0,0,0.5);
    }}
    div[data-testid="stVerticalBlockBorderWrapper"]:hover::before,
    [data-testid="stLayoutWrapper"] > [data-testid="stVerticalBlock"]:not([class*="st-key-"]):hover::before {{
        background: linear-gradient(90deg, transparent, {hex_to_rgba(COLORS['primary_deep'], 0.35)}, transparent);
    }}

    .kpi-icon {{
        display:inline-flex; align-items:center; justify-content:center;
        width:28px; height:28px; border-radius:50%; font-size:15px; margin-right:8px;
        box-shadow: 0 0 0 1px {hex_to_rgba(COLORS['primary'], 0.15)};
    }}
    .kpi-label {{
        color:{COLORS['text_soft']}; font-size:0.72rem; font-weight:600;
        text-transform: uppercase; letter-spacing: 0.15em; font-family: 'Inter', sans-serif;
    }}
    .kpi-value {{
        font-size:2.3rem; font-weight:400; margin: 4px 0 2px 0; letter-spacing:0.01em;
        font-family: 'Bebas Neue', sans-serif !important;
    }}
    .pill {{
        display:inline-block; padding:3px 11px; border-radius:20px;
        font-size:0.76rem; font-weight:700;
    }}
    .pill-up-bad {{ background-color: {hex_to_rgba(COLORS['danger'], 0.20)}; color:{COLORS['danger']} !important; }}
    .pill-up-good {{ background-color: {hex_to_rgba(COLORS['success'], 0.18)}; color:{COLORS['success']} !important; }}
    .pill-down-good {{ background-color: {hex_to_rgba(COLORS['success'], 0.18)}; color:{COLORS['success']} !important; }}
    .pill-flat {{ background-color: {hex_to_rgba(COLORS['text_soft'], 0.15)}; color:{COLORS['text_soft']} !important; }}
    .pill-ok {{ background-color: {hex_to_rgba(COLORS['success'], 0.18)}; color:{COLORS['success']} !important; }}
    .pill-warn {{ background-color: {hex_to_rgba(COLORS['warning'], 0.20)}; color:{COLORS['warning']} !important; }}
    .pill-risk {{ background-color: {hex_to_rgba(COLORS['danger'], 0.22)}; color:{COLORS['danger']} !important; }}

    /* Category panels. A quarter-width column is too narrow for a
       label-left/value-right row — long labels and values collide and wrap
       mid-string ("AED / 42,000/mo"). Stacking the label over the value
       gives each the full card width, so nothing breaks awkwardly, and it
       reuses the label-over-value pattern the rest of the system already
       uses for figures. */
    /* Stretch every card to the tallest in the row so the bottoms line up.
       The height has to be carried all the way down the column -> wrapper ->
       block chain, or the innermost card still collapses to its content.
       Both the old and current Streamlit container test-ids are listed so
       this keeps working across versions. */
    .st-key-gd_catpanels [data-testid="stColumn"] {{ display: flex; }}
    .st-key-gd_catpanels [data-testid="stColumn"] > div,
    .st-key-gd_catpanels [data-testid="stLayoutWrapper"],
    .st-key-gd_catpanels [data-testid="stLayoutWrapper"] > [data-testid="stVerticalBlock"],
    .st-key-gd_catpanels div[data-testid="stVerticalBlockBorderWrapper"] {{
        width: 100%; height: 100%;
    }}
    .cat-panel {{ padding: 14px 16px 6px 16px; }}
    .cat-panel h4 {{ margin: 0 0 12px 0 !important; font-size: 0.95rem !important; }}
    .cat-metric {{ padding: 9px 0; border-bottom: 1px solid {COLORS['border']}; }}
    .cat-metric:last-child {{ border-bottom: none; }}
    .cat-metric-label {{
        color: {COLORS['text_soft']} !important;
        font-size: 0.64rem; font-weight: 600;
        text-transform: uppercase; letter-spacing: 0.13em;
        line-height: 1.35;
    }}
    .cat-metric-value {{
        color: {COLORS['text']} !important;
        font-size: 0.98rem; font-weight: 600; line-height: 1.25;
        margin-top: 4px;
    }}

    .info-row {{
        padding:8px 0; border-bottom:1px solid {COLORS['border']};
        font-size:0.92rem; line-height:1.5; display:flex; gap:8px;
    }}
    .info-row:last-child {{ border-bottom:none; }}
    .info-row::before {{ content:"—"; color:{COLORS['primary']}; flex-shrink:0; }}

    .hero-badge {{
        display:inline-block; padding:5px 14px; border-radius:20px; font-size:0.78rem;
        font-weight:700; background-color:{COLORS['surface']}; border:1px solid {COLORS['border']};
        color:{COLORS['text_soft']} !important; margin-right:8px;
    }}

    /* Underline-style tabs — matches the DESIGN.md's nav spec: minimalist
       uppercase text links, active state = 2px gold underline, no fill. */
    .stTabs [data-baseweb="tab-list"] {{
        background-color: transparent; border-bottom: 1px solid {COLORS['border']};
        gap: 24px; padding: 0 0 0 2px;
    }}
    .stTabs [data-baseweb="tab"] {{
        color: {COLORS['text_soft']} !important; border-radius: 0 !important;
        padding: 8px 2px !important; text-transform: uppercase;
        font-size: 0.78rem; letter-spacing: 0.06em; font-weight: 600;
        background-color: transparent !important;
    }}
    .stTabs [aria-selected="true"] {{
        background-color: transparent !important;
        color: {COLORS['primary']} !important; font-weight:700 !important;
        border-bottom: 2px solid {COLORS['primary']} !important;
    }}
    .stTabs [data-baseweb="tab-highlight"] {{ background-color: transparent !important; }}
    .stTabs [data-baseweb="tab-border"] {{ display:none !important; }}

    div[data-testid="stMetric"] {{ background-color: transparent; }}

    /* Sliders / inputs on dark */
    div[data-baseweb="slider"] div[role="slider"] {{ background-color: {COLORS['primary']} !important; }}
    .stSlider [data-testid="stTickBar"] {{ display:none; }}
</style>
""", unsafe_allow_html=True)

PLOTLY_DARK = dict(
    paper_bgcolor="rgba(0,0,0,0)",
    plot_bgcolor="rgba(0,0,0,0)",
    font=dict(color=COLORS["text"], size=12),
    margin=dict(l=10, r=10, t=25, b=10),
)
GRID_COLOR = "rgba(255,255,255,0.06)"

# ----------------------------------------------------------------------
# DISPLAY-PRECISION FORMATTERS (§A9) — the only place a labour_calc figure
# gets rounded. Every value from labour_calc.py stays full-float until it
# reaches one of these; `None` (an undefined ratio, or missing revenue)
# always renders as "—", never 0.00 or 0%.
# ----------------------------------------------------------------------
def fmt_hours(value):
    return "—" if value is None else f"{value:.2f}"

def fmt_pct(value):
    return "—" if value is None else f"{value:.1f}%"

def fmt_aed(value):
    return "—" if value is None else f"AED {value:,.2f}"

def fmt_ratio(value):
    return "—" if value is None else f"{value:.2f}×"

def sparkline(values, color):
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        y=values, mode="lines", line=dict(color=color, width=2.5),
        fill="tozeroy", fillcolor=color.replace(")", ",0.12)").replace("rgb", "rgba") if "rgb" in color else None,
    ))
    fig.update_layout(
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        margin=dict(l=0, r=0, t=2, b=2), height=48,
        xaxis=dict(visible=False), yaxis=dict(visible=False), showlegend=False,
    )
    return fig

# ----------------------------------------------------------------------
# SAMPLE / BENCHMARK DATA (placeholder until mentor-provided actuals)
# ----------------------------------------------------------------------
months = ["Feb", "Mar", "Apr", "May", "Jun", "Jul"]
food_cost_trend = [29.8, 30.1, 30.6, 30.9, 31.0, 31.4]
target_line = [30.0] * 6
wastage_trend = [5.4, 5.2, 5.1, 4.9, 4.9, 4.8]
margin_trend = [42.3, 42.0, 41.8, 42.0, 42.2, 42.1]
cost_unit_trend = [3.55, 3.62, 3.68, 3.74, 3.80, 3.85]

baseline = {
    "food_cost_pct": 31.4,
    "wastage_pct": 4.8,
    "gross_margin_pct": 42.1,
    "cost_per_unit": 3.85,
    "standard_cost_per_unit": 3.60,
}

category_panels = {
    "🥖 Production": {
        "Batch efficiency": "88%",
        "Labour cost per unit": "AED 1.10",
        "Actual vs standard cost": "+6.9%",
    },
    "🚚 Procurement": {
        "Purchase price variance": "+3.2%",
        "Average supplier lead time": "2.4 days",
        "Top-3 flour supplier share": "76%",
    },
    "📦 Warehouse & inventory": {
        "Inventory turnover": "14.2x per year",
        "Expired / damaged stock": "AED 8,400",
        "Slow-moving SKUs": "6 lines",
    },
    "💰 Cost control & margin": {
        "Estimated margin leakage": "AED 42,000 / month",
        "Packaging cost": "2.1% of sales",
        "Overhead allocation": "9.4% of sales",
    },
}

# ----------------------------------------------------------------------
# STATUS REPORT (EXCEL) + EMAIL DELIVERY
# Each "Email this ..." button builds a workbook tailored to the section
# it's called from, not a generic dashboard dump — Performance Tracker,
# Scenario & Resilience, a department comparison, or the Company Profile.
# ----------------------------------------------------------------------
XL_PRIMARY = "B9A0DC"
XL_DARK = "1A1424"
XL_HEADER_FONT = "FFFFFF"
XL_TEXT = "2B2733"
XL_SOFT_TEXT = "6B6577"
XL_LIGHT_FILL = "F3EFFA"
XL_BORDER_COLOR = "D9D2E9"

def _xl_border():
    side = Side(style="thin", color=XL_BORDER_COLOR)
    return Border(left=side, right=side, top=side, bottom=side)

def _xl_title(ws, title, subtitle=None):
    ws["A1"] = title
    ws["A1"].font = Font(name="Arial", size=16, bold=True, color=XL_DARK)
    row = 2
    if subtitle:
        ws.cell(row=2, column=1, value=subtitle).font = Font(name="Arial", size=10, italic=True, color=XL_SOFT_TEXT)
        row = 3
    return row + 1

def _xl_section(ws, row, text):
    ws.cell(row=row, column=1, value=text).font = Font(name="Arial", size=12, bold=True, color=XL_DARK)
    return row + 1

def _xl_header_row(ws, row, headers, start_col=1):
    border = _xl_border()
    for i, h in enumerate(headers):
        c = ws.cell(row=row, column=start_col + i, value=h)
        c.font = Font(name="Arial", size=10, bold=True, color=XL_HEADER_FONT)
        c.fill = PatternFill("solid", fgColor=XL_PRIMARY)
        c.alignment = Alignment(horizontal="left", vertical="center")
        c.border = border
    return row + 1

def _xl_row(ws, row, values, start_col=1, formats=None, zebra=False):
    border = _xl_border()
    fill = PatternFill("solid", fgColor=XL_LIGHT_FILL) if zebra else None
    for i, v in enumerate(values):
        c = ws.cell(row=row, column=start_col + i, value=v)
        c.font = Font(name="Arial", size=10, color=XL_TEXT)
        c.border = border
        if fill:
            c.fill = fill
        if formats and i < len(formats) and formats[i]:
            c.number_format = formats[i]
    return row + 1

def _xl_autosize(ws, max_width=64):
    widths = {}
    for row in ws.iter_rows():
        for cell in row:
            if cell.value is None or hasattr(cell, "column_letter") is False:
                continue
            try:
                col = cell.column_letter
            except AttributeError:
                continue
            widths[col] = max(widths.get(col, 8), min(len(str(cell.value)) + 3, max_width))
    for col, w in widths.items():
        ws.column_dimensions[col].width = w

def _xl_footer(ws, row):
    row += 1
    ws.cell(row=row, column=1,
             value=f"Generated {datetime.now().strftime('%d %B %Y, %H:%M')} · GIP III · Bakery division only, catering excluded"
             ).font = Font(name="Arial", size=8, italic=True, color=XL_SOFT_TEXT)
    row += 1
    ws.cell(row=row, column=1,
             value="Prepared by Parineeta Jain, Rajveer Singh, and Tarang Gupta · Figures are illustrative "
                   "benchmarks pending Grandiose-provided actuals."
             ).font = Font(name="Arial", size=8, italic=True, color=XL_SOFT_TEXT)
    return row

PCT_FMT = '0.0"%"'
AED_FMT = '"AED "#,##0.00'
NUM_FMT = "#,##0"


def _build_performance_tracker_sheet(wb):
    ws = wb.create_sheet("Performance Tracker")
    row = _xl_title(ws, "Grandiose Bakery — Performance Tracker",
                     "Where the bakery division stands today")

    row = _xl_section(ws, row, "Headline KPIs")
    row = _xl_header_row(ws, row, ["Metric", "Value", "vs. target / last month"])
    kpi_rows = [
        ("Food cost %", baseline["food_cost_pct"], "1.2pt above 30% target"),
        ("Wastage %", baseline["wastage_pct"], "Down 0.3pt vs last month"),
        ("Gross margin %", baseline["gross_margin_pct"], "Flat vs last month"),
        ("Cost per unit (AED)", baseline["cost_per_unit"], f"Standard: AED {baseline['standard_cost_per_unit']:.2f}"),
    ]
    for i, (label, val, note) in enumerate(kpi_rows):
        fmt = AED_FMT if "unit" in label.lower() else PCT_FMT
        row = _xl_row(ws, row, [label, val, note], formats=[None, fmt, None], zebra=(i % 2 == 0))
    row += 1

    row = _xl_section(ws, row, "Food cost % — last 6 months vs. 30% target")
    row = _xl_header_row(ws, row, ["Month", "Food cost %", "Target %"])
    for i, (m, fc, tg) in enumerate(zip(months, food_cost_trend, target_line)):
        row = _xl_row(ws, row, [m, fc, tg], formats=[None, PCT_FMT, PCT_FMT], zebra=(i % 2 == 0))
    row += 1

    row = _xl_section(ws, row, "Category panels")
    row = _xl_header_row(ws, row, ["Category", "Metric", "Value"])
    zebra = False
    for cat_name, metrics in category_panels.items():
        clean_cat = cat_name.split(" ", 1)[-1]
        for k, v in metrics.items():
            row = _xl_row(ws, row, [clean_cat, k, v], zebra=zebra)
            zebra = not zebra
    row += 1

    _xl_footer(ws, row)
    _xl_autosize(ws)
    return ws


def _build_scenario_resilience_sheets(wb, context):
    infl = context.get("inflation", {})
    disr = context.get("disruption", {})
    pand = context.get("pandemic", {})
    supp = context.get("supplier", {})

    # --- Inflation sensitivity ---
    ws = wb.create_sheet("Inflation Sensitivity")
    row = _xl_title(ws, "Scenario & Resilience — Inflation Sensitivity",
                     "Dual-track inflation input, net of any subsidy/offset")
    row = _xl_section(ws, row, "Assumptions used")
    row = _xl_header_row(ws, row, ["Input", "Value"])
    row = _xl_row(ws, row, ["Headline inflation forecast", infl.get("headline_inf", 0)], formats=[None, PCT_FMT], zebra=True)
    row = _xl_row(ws, row, ["Actual food-input inflation", infl.get("food_inf", 0)], formats=[None, PCT_FMT])
    row = _xl_row(ws, row, ["Subsidy / price-cap offset (AED/unit)", infl.get("subsidy_offset", 0)], formats=[None, AED_FMT], zebra=True)
    row += 1
    row = _xl_section(ws, row, "Resulting cost per unit")
    row = _xl_header_row(ws, row, ["Scenario", "Cost per unit (AED)"])
    row = _xl_row(ws, row, ["Current", infl.get("base_cost", 0)], formats=[None, AED_FMT], zebra=True)
    row = _xl_row(ws, row, ["Headline scenario", infl.get("adj_cost_headline", 0)], formats=[None, AED_FMT])
    row = _xl_row(ws, row, ["Food-inflation scenario", infl.get("adj_cost_food", 0)], formats=[None, AED_FMT], zebra=True)
    row = _xl_row(ws, row, ["Food cost % under food-inflation scenario", infl.get("food_cost_pct_adj", 0)], formats=[None, PCT_FMT])
    _xl_footer(ws, row)
    _xl_autosize(ws)

    # --- Supply disruption ---
    ws2 = wb.create_sheet("Supply Disruption")
    row = _xl_title(ws2, "Scenario & Resilience — Supply Disruption Risk",
                     "UAE imports 80-90% of its food, so this is mainly an import-shock model")
    row = _xl_section(ws2, row, "Selected scenario")
    row = _xl_header_row(ws2, row, ["Field", "Value"])
    row = _xl_row(ws2, row, ["Scenario", disr.get("scenario", "—")], zebra=True)
    row = _xl_row(ws2, row, ["Lead-time extension (days)", disr.get("delay_days", 0)], formats=[None, NUM_FMT])
    row = _xl_row(ws2, row, ["Cost premium %", disr.get("cost_premium", 0)], formats=[None, PCT_FMT], zebra=True)
    row = _xl_row(ws2, row, ["Stockout probability", disr.get("stockout_prob", 0) * 100], formats=[None, PCT_FMT])
    row += 1
    row = _xl_section(ws2, row, "Buffer vs. expected stockout cost")
    row = _xl_header_row(ws2, row, ["Metric", "AED"])
    row = _xl_row(ws2, row, ["Cost of holding buffer stock", disr.get("buffer_stock_cost", 0)], formats=[None, AED_FMT], zebra=True)
    row = _xl_row(ws2, row, ["Expected cost of stockout", disr.get("expected_stockout_cost", 0)], formats=[None, AED_FMT])
    _xl_footer(ws2, row)
    _xl_autosize(ws2)

    # --- Pandemic preparedness ---
    ws3 = wb.create_sheet("Pandemic Preparedness")
    row = _xl_title(ws3, "Scenario & Resilience — Pandemic Preparedness",
                     "Demand-side and supply-side resilience, tracked together")
    row = _xl_section(ws3, row, "Customer retention / attraction")
    row = _xl_header_row(ws3, row, ["Metric", "Value"])
    row = _xl_row(ws3, row, ["Repeat-purchase rate", pand.get("repeat_rate", 0)], formats=[None, PCT_FMT], zebra=True)
    row = _xl_row(ws3, row, ["Delivery/online-order revenue share", pand.get("delivery_share", 0)], formats=[None, PCT_FMT])
    row = _xl_row(ws3, row, ["Average basket size (AED)", pand.get("basket_size", 0)], formats=[None, AED_FMT], zebra=True)
    row += 1
    row = _xl_section(ws3, row, "Supply chain optimization")
    row = _xl_header_row(ws3, row, ["Metric", "Value"])
    row = _xl_row(ws3, row, ["Safety stock (days of cover)", pand.get("safety_days", 0)], formats=[None, NUM_FMT], zebra=True)
    row = _xl_row(ws3, row, ["Qualified alternate suppliers / key ingredient", pand.get("alt_suppliers", 0)], formats=[None, NUM_FMT])
    row = _xl_row(ws3, row, ["% ingredients single-sourced", pand.get("single_sourced", 0)], formats=[None, PCT_FMT], zebra=True)
    _xl_footer(ws3, row)
    _xl_autosize(ws3)

    # --- Supplier concentration ---
    ws4 = wb.create_sheet("Supplier Concentration")
    row = _xl_title(ws4, "Scenario & Resilience — Supplier Concentration",
                     "HHI-style concentration index, following Ali et al. (2022)'s UAE cereal-supply-risk methodology")
    row = _xl_section(ws4, row, "Concentration index")
    row = _xl_header_row(ws4, row, ["Metric", "Value"])
    row = _xl_row(ws4, row, ["HHI", supp.get("hhi", 0)], formats=[None, NUM_FMT], zebra=True)
    row = _xl_row(ws4, row, ["Risk level", supp.get("risk_label", "—")])
    row += 1
    shares_df = supp.get("shares_df")
    if shares_df is not None and not shares_df.empty:
        row = _xl_section(ws4, row, "Supplier spend mix")
        row = _xl_header_row(ws4, row, list(shares_df.columns))
        for i, (_, r) in enumerate(shares_df.iterrows()):
            row = _xl_row(ws4, row, list(r.values), zebra=(i % 2 == 0))
        row += 1
    alt_df = supp.get("alt_df")
    if alt_df is not None and not alt_df.empty:
        row = _xl_section(ws4, row, "Pre-identified alternate suppliers")
        row = _xl_header_row(ws4, row, list(alt_df.columns))
        for i, (_, r) in enumerate(alt_df.iterrows()):
            row = _xl_row(ws4, row, list(r.values), zebra=(i % 2 == 0))
    _xl_footer(ws4, row)
    _xl_autosize(ws4)


def _build_department_sheet(wb, context):
    dept_name = context.get("department", "Department")
    ws = wb.create_sheet(dept_name[:31] if dept_name else "Department")
    row = _xl_title(ws, f"Employee Portal — {dept_name} Comparison",
                     "Performance ranking within this department")
    headcount = context.get("headcount", "—")
    tag = context.get("tag_label")
    row = _xl_section(ws, row, f"Headcount on this line: {headcount}" + (f" · {tag}" if tag else ""))
    df = context.get("summary_df")
    if df is not None and not df.empty:
        row = _xl_header_row(ws, row, list(df.columns))
        for i, (_, r) in enumerate(df.iterrows()):
            row = _xl_row(ws, row, list(r.values), zebra=(i % 2 == 0))
    _xl_footer(ws, row)
    _xl_autosize(ws)
    return ws


def _build_company_profile_sheet(wb, context):
    ws = wb.create_sheet("Company Profile")
    row = _xl_title(ws, "Grandiose Bakery — Company Profile",
                     "Everything collected from the 27 Jul meeting with the GM")
    company_info = context.get("company_info", [])
    for icon, title, points in company_info:
        row = _xl_section(ws, row, f"{icon} {title}")
        for i, p in enumerate(points):
            row = _xl_row(ws, row, [p], zebra=(i % 2 == 0))
        row += 1
    _xl_footer(ws, row)
    _xl_autosize(ws)
    return ws


def _coerce_cell(value):
    """Try to turn AI-returned strings that look numeric into real numbers,
    so Excel treats them as numbers (right-aligned, sortable) not text.
    Leaves percentage strings ('5.0%') and currency strings ('AED 3.85')
    as readable text rather than stripping their meaning away."""
    if isinstance(value, (int, float)):
        return value
    if isinstance(value, str):
        stripped = value.strip()
        if "%" in stripped or "AED" in stripped:
            return stripped
        cleaned = stripped.replace(",", "")
        try:
            if "." in cleaned:
                return float(cleaned)
            return int(cleaned)
        except (ValueError, TypeError):
            return value
    return value


def _build_smart_upload_sheets(wb, context):
    result = context.get("result", {})
    source_files = context.get("source_files", [])

    ws = wb.create_sheet("Summary")
    row = _xl_title(ws, "Grandiose Bakery — Processed Data Report",
                     "AI-interpreted from your uploaded file(s)")
    row = _xl_section(ws, row, "Source files")
    row = _xl_header_row(ws, row, ["File name"])
    for i, fname in enumerate(source_files):
        row = _xl_row(ws, row, [fname], zebra=(i % 2 == 0))
    row += 1
    row = _xl_section(ws, row, "Detected data type")
    ws.cell(row=row, column=1, value=result.get("detected_data_type", "—")).font = Font(name="Arial", size=10, color=XL_TEXT)
    row += 2
    row = _xl_section(ws, row, "Overall summary")
    for line in (result.get("summary", "") or "").split("\n"):
        if line.strip():
            ws.cell(row=row, column=1, value=line.strip()).font = Font(name="Arial", size=10, color=XL_TEXT)
            row += 1
    _xl_footer(ws, row)
    _xl_autosize(ws)

    for sheet_info in result.get("sheets", []):
        raw_name = sheet_info.get("sheet_name") or sheet_info.get("title") or "Sheet"
        safe_name = re.sub(r'[\\/*?:\[\]]', "-", raw_name)[:31] or "Sheet"
        base_name, n = safe_name, 1
        while safe_name in wb.sheetnames:
            n += 1
            safe_name = f"{base_name[:28]}_{n}"
        ws2 = wb.create_sheet(safe_name)
        row = _xl_title(ws2, sheet_info.get("title", raw_name))
        columns = sheet_info.get("columns", [])
        data_rows = sheet_info.get("rows", [])
        if columns:
            row = _xl_header_row(ws2, row, columns)
            for i, r in enumerate(data_rows):
                r = list(r) + [""] * (len(columns) - len(r)) if len(r) < len(columns) else r[:len(columns)]
                row = _xl_row(ws2, row, [_coerce_cell(v) for v in r], zebra=(i % 2 == 0))
            row += 1
        insights = sheet_info.get("insights", [])
        if insights:
            row = _xl_section(ws2, row, "Insights")
            for i, ins in enumerate(insights):
                row = _xl_row(ws2, row, [ins], zebra=(i % 2 == 0))
        _xl_footer(ws2, row)
        _xl_autosize(ws2)


def extract_pdf_text(file_bytes):
    try:
        reader = PdfReader(BytesIO(file_bytes))
        pages = [p.extract_text() or "" for p in reader.pages[:30]]
        return "\n".join(pages)
    except Exception as e:
        return f"[Could not read PDF: {e}]"


def extract_docx_text(file_bytes):
    try:
        doc = DocxDocument(BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for r in table.rows:
                parts.append(" | ".join(c.text for c in r.cells))
        return "\n".join(parts)
    except Exception as e:
        return f"[Could not read DOCX: {e}]"


def extract_excel_text(file_bytes, filename):
    try:
        engine = "openpyxl" if filename.lower().endswith((".xlsx", ".xlsm")) else None
        sheets = pd.read_excel(BytesIO(file_bytes), sheet_name=None, engine=engine) if not filename.lower().endswith(".csv") \
            else {"Sheet1": pd.read_csv(BytesIO(file_bytes))}
        parts = []
        for name, df in sheets.items():
            parts.append(f"--- Sheet: {name} ---")
            parts.append(df.to_csv(index=False))
        return "\n".join(parts)
    except Exception as e:
        return f"[Could not read spreadsheet: {e}]"


MAX_UPLOAD_CHARS = 60000  # keeps API cost/latency bounded across a batch of files

@st.cache_resource
def get_anthropic_client():
    try:
        api_key = st.secrets["ANTHROPIC_API_KEY"]
    except (KeyError, FileNotFoundError):
        return None
    try:
        return Anthropic(api_key=api_key)
    except Exception:
        return None


INTERPRETER_SYSTEM_PROMPT = """You are a financial and operations analyst working on a cost-optimization \
dashboard for Grandiose Bakery, a UAE bakery division (GIP III project). You will be given raw, rough, \
messy content extracted from one or more uploaded files (Excel, PDF, or Word) — it could be production logs, \
procurement records, wastage logs, sales data, inventory data, employee shift logs, or something else entirely.

Your job:
1. Identify what kind of business data this actually is.
2. Clean and interpret the rough numbers — infer column meaning even if headers are messy, missing, or \
inconsistent. Handle obvious typos and unit inconsistencies sensibly.
3. Where the data supports it, compute relevant KPIs using standard bakery cost-optimization metrics: \
wastage % (wasted / (produced + wasted) * 100), cost per unit, batch-time adherence %, quality pass rate, \
supplier concentration (HHI), productivity (value generated per employee per day), or straightforward \
sums/averages/trends — but ONLY compute what the data actually supports. Never fabricate numbers that \
aren't derivable from what was given.
4. Write 2-4 sentences of plain-English, management-readable insight per output sheet — what the numbers \
show and what it implies, not just a restatement of the table.

Respond with STRICT JSON only, no markdown fences, no commentary outside the JSON, matching exactly this \
shape:
{
  "detected_data_type": "short description of what this data is",
  "summary": "2-4 sentence overall summary of what was found across all files",
  "sheets": [
    {
      "sheet_name": "short sheet name, max 31 chars, no slashes or brackets",
      "title": "human-readable title for this sheet",
      "columns": ["Column A", "Column B", ...],
      "rows": [["value", "value", ...], ...],
      "insights": ["insight sentence 1", "insight sentence 2"]
    }
  ]
}

Produce 1-4 sheets depending on what's genuinely present in the data — don't invent sheets with no real \
content. If the uploaded content is unreadable or empty, still return valid JSON with an empty "sheets" \
list and explain why in "summary"."""


def interpret_uploaded_files(file_contents: dict):
    """file_contents: {filename: extracted_text}. Returns (result_dict, error_message)."""
    client = get_anthropic_client()
    if client is None:
        return None, ("AI processing isn't configured yet. Add ANTHROPIC_API_KEY in Streamlit Cloud's "
                       "Secrets settings (App settings -> Secrets) — never in the code or GitHub repo.")

    combined = ""
    for fname, text in file_contents.items():
        combined += f"\n\n===== FILE: {fname} =====\n{text}"
    truncated = len(combined) > MAX_UPLOAD_CHARS
    if truncated:
        combined = combined[:MAX_UPLOAD_CHARS]

    try:
        response = client.messages.create(
            model="claude-sonnet-5",
            max_tokens=8000,
            system=INTERPRETER_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": f"Here is the extracted content to interpret:{combined}"}],
        )
        raw_text = "".join(block.text for block in response.content if hasattr(block, "text"))
        cleaned = raw_text.strip()
        if cleaned.startswith("```"):
            cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", cleaned.strip())
        result = json.loads(cleaned)
        if truncated:
            result["summary"] = (result.get("summary", "") +
                                  " (Note: input was long and was truncated before analysis — "
                                  "results reflect only the portion processed.)")
        return result, None
    except json.JSONDecodeError:
        return None, "The AI response couldn't be parsed as structured data. Try again, or with a smaller/simpler file."
    except Exception as e:
        return None, f"Could not process files: {e}"


def build_excel_report(section, context=None):
    """Builds a section-specific, formatted Excel workbook and returns raw bytes."""
    context = context or {}
    wb = Workbook()
    wb.remove(wb.active)

    if section == "scenario_resilience":
        _build_scenario_resilience_sheets(wb, context)
    elif section == "employee_department":
        _build_department_sheet(wb, context)
    elif section == "company_profile":
        _build_company_profile_sheet(wb, context)
    elif section == "smart_upload":
        _build_smart_upload_sheets(wb, context)
    else:
        _build_performance_tracker_sheet(wb)

    bio = BytesIO()
    wb.save(bio)
    return bio.getvalue()


SECTION_LABELS = {
    "performance_tracker": ("Performance Tracker", "Performance_Tracker"),
    "scenario_resilience": ("Scenario & Resilience", "Scenario_Resilience"),
    "employee_department": ("Employee Department Comparison", "Department_Comparison"),
    "company_profile": ("Company Profile", "Company_Profile"),
    "smart_upload": ("Processed Data Report", "Processed_Data"),
}

def send_report_email(recipient_email, note="", section="performance_tracker", context=None):
    """Sends a section-specific Excel report to recipient_email using SMTP
    credentials stored in Streamlit secrets. Returns (success: bool, message: str)."""
    try:
        sender_email = st.secrets["EMAIL_ADDRESS"]
        sender_password = st.secrets["EMAIL_APP_PASSWORD"]
    except (KeyError, FileNotFoundError):
        return False, ("Email is not configured yet. Add EMAIL_ADDRESS and EMAIL_APP_PASSWORD "
                        "in Streamlit Cloud's Secrets settings (App settings -> Secrets) — "
                        "never in the code or GitHub repo.")

    smtp_server = st.secrets.get("SMTP_SERVER", "smtp.gmail.com")
    smtp_port = int(st.secrets.get("SMTP_PORT", 587))
    section_title, section_slug = SECTION_LABELS.get(section, SECTION_LABELS["performance_tracker"])

    msg = MIMEMultipart()
    msg["From"] = sender_email
    msg["To"] = recipient_email
    msg["Subject"] = f"Grandiose Bakery — {section_title} Report"
    body = (f"Hi,\n\nPlease find attached the {section_title} report from the Grandiose bakery "
            f"dashboard, generated just now.\n")
    if note:
        body += f"\nNote from sender:\n{note}\n"
    body += "\n— Sent automatically from the Grandiose Bakery Dashboard (GIP III)"
    msg.attach(MIMEText(body, "plain"))

    xlsx_bytes = build_excel_report(section, context)
    attachment = MIMEApplication(
        xlsx_bytes, _subtype="vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    filename = f"Grandiose_Bakery_{section_slug}_{datetime.now().strftime('%Y%m%d')}.xlsx"
    attachment.add_header("Content-Disposition", "attachment", filename=filename)
    msg.attach(attachment)

    try:
        with smtplib.SMTP(smtp_server, smtp_port, timeout=15) as server:
            server.starttls()
            server.login(sender_email, sender_password)
            server.send_message(msg)
        return True, f"Report sent to {recipient_email}."
    except Exception as e:
        return False, f"Could not send email: {e}"


EMAIL_PATTERN = r"^[^@\s]+@[^@\s]+\.[^@\s]+$"

def render_email_share(key_suffix, label="Email this report", section="performance_tracker", context=None):
    """Small 📧 popover button — usable right under any chart/section so
    people don't have to go back to the sidebar to share the dashboard.
    Sends a formatted Excel workbook tailored to the calling section."""
    section_title, _ = SECTION_LABELS.get(section, SECTION_LABELS["performance_tracker"])
    with st.popover("📧", help=label):
        st.markdown(f"**{label}**")
        recipient = st.text_input("Recipient email", placeholder="name@company.com", key=f"email_recipient_{key_suffix}")
        note = st.text_area("Add a note (optional)", key=f"email_note_{key_suffix}", height=68)
        if st.button("Send report", key=f"send_email_btn_{key_suffix}", width='stretch', type="primary"):
            if not recipient or not re.match(EMAIL_PATTERN, recipient):
                st.error("Enter a valid email address.")
            else:
                with st.spinner("Building and sending Excel report..."):
                    success, message = send_report_email(recipient, note, section=section, context=context)
                if success:
                    st.success(message)
                else:
                    st.error(message)
        st.caption(f"Sends a formatted Excel workbook of the current {section_title} data.")


st.sidebar.markdown(f"""
<div style='background-color:{COLORS['surface']}; border:1px solid {COLORS['border']};
     border-radius:4px; padding:16px; margin-bottom:14px;'>
  <div style="font-family:'Bebas Neue',sans-serif; font-size:1.9rem; color:{COLORS['primary']};
       text-transform:uppercase; letter-spacing:0.02em; line-height:1;">Grandiose</div>
  <div style='color:{COLORS['text_soft']}; font-size:0.8rem; margin-top:8px;'>
    GIP III — Financial performance &amp; cost optimization<br>Bakery division only · catering excluded
  </div>
</div>
""", unsafe_allow_html=True)

st.sidebar.markdown(f"""
<div style='background-color:{COLORS['surface']}; border:1px solid {COLORS['border']};
     border-radius:4px; padding:14px; margin-bottom:14px;'>
  <div style='color:{COLORS['text_soft']}; font-size:0.78rem; font-weight:700; letter-spacing:0.5px;'>PREPARED BY</div>
  <div style='margin-top:6px; font-size:0.9rem;'>Parineeta Jain</div>
  <div style='font-size:0.9rem;'>Rajveer Singh</div>
  <div style='font-size:0.9rem;'>Tarang Gupta</div>
</div>
""", unsafe_allow_html=True)

st.sidebar.markdown(f"""
<div style='background-color:{hex_to_rgba(COLORS['primary'], 0.1)}; border:1px solid {hex_to_rgba(COLORS['primary'], 0.3)};
     border-radius:4px; padding:14px; font-size:0.82rem; color:{COLORS['text']} !important;'>
  ⚠️ All figures shown are illustrative benchmarks pending Grandiose's actual bakery data.
</div>
""", unsafe_allow_html=True)

with st.sidebar.expander("📧  Email this report", expanded=False):
    recipient = st.text_input("Recipient email", placeholder="name@company.com", key="email_recipient")
    note = st.text_area("Add a note (optional)", key="email_note", height=70)
    if st.button("Send report", key="send_email_btn", width='stretch', type="primary"):
        email_pattern = r"^[^@\s]+@[^@\s]+\.[^@\s]+$"
        if not recipient or not re.match(email_pattern, recipient):
            st.error("Enter a valid email address.")
        else:
            with st.spinner("Building and sending Excel report..."):
                success, message = send_report_email(recipient, note, section="performance_tracker")
            if success:
                st.success(message)
            else:
                st.error(message)
    st.caption("Sends a formatted Excel workbook of the current Performance Tracker data.")


# ----------------------------------------------------------------------
# EMPLOYEE PORTAL — data layer
# Uses Supabase (Postgres) when configured via secrets; falls back to an
# in-session store so the portal is still fully clickable/demoable before
# Supabase is wired up. Data in fallback mode does not persist across
# app restarts — see README for the one-time Supabase setup.
# ----------------------------------------------------------------------

# Actual staff breakdown by line/department (approximate headcounts).
DEPARTMENT_HEADCOUNT = {
    "Baklava": 10,
    "French Bakery": 10,
    "Arabic Bread": 7,
    "Viennoiserie (Croissant/Danish)": 7,
    "Tahina": 5,
    "QC": 3,
    "Packing & Dispatch": 15,
    "Store": 2,
    "Maintenance": 4,
    "GM": 1,
    "Office": 10,
    "Drivers": 7,
    "Third-Party Cleaners": 10,
}

# Company-wide headcount context from the 27 Jul MoM. Note this figure (~110)
# is larger than the sum of the itemized bakery lines above (91) — likely
# because it includes catering/other roles not broken out per line. Shown
# for context only; the itemized breakdown above remains the bakery-scope
# source of truth for GIP III.
COMPANY_HEADCOUNT_CURRENT = 110
COMPANY_HEADCOUNT_TARGET = 200

# Illustrative daily capacity targets (units/day) for production lines only —
# PLACEHOLDER figures pending Grandiose's actual max-output numbers per line.
# The MoM gives an overall utilization figure (30-35%, target 100% within a
# year) but not a per-line breakdown, so utilization % shown against these
# targets should be treated as directional, not precise, until replaced.
DEPARTMENT_CAPACITY_TARGET = {
    "Baklava": 400,
    "French Bakery": 500,
    "Arabic Bread": 900,
    "Viennoiserie (Croissant/Danish)": 450,
    "Tahina": 250,
}

# Lines the MoM specifically called out — used to tag department cards/tables.
DEPARTMENT_PRIORITY_TAG = {
    "Arabic Bread": ("🔥 Top revenue driver", "pill-ok"),
    "Tahina": ("⚠️ Ramp-up focus (underutilized)", "pill-warn"),
    "Baklava": ("⚠️ Ramp-up focus (underutilized)", "pill-warn"),
}

# GM's stated productivity benchmark: each employee = AED 1,000/day value.
VALUE_PER_EMPLOYEE_DAY_TARGET = 1000
# Standard shift length per the MoM ("Standard shift is 9 hours, including a
# break") — used to normalize revenue/hour up to a full standard day.
STANDARD_SHIFT_HOURS = 9.0

# Company-wide wastage target from the MoM (current is 2-3%, target is 1%).
WASTAGE_TARGET_PCT = 1.0

# Production lines vs. support functions — labour_calc's §A7 revenue
# attribution treats these differently: production departments log revenue
# directly, support departments receive an allocated share instead.
PRODUCTION_DEPARTMENTS = set(DEPARTMENT_CAPACITY_TARGET)

# Support department's share of production-line revenue, by that
# department's share of total support headcount. PLACEHOLDER basis pending
# real figures from finance (§A7) — TODO: CONFIRM WITH GM/FINANCE.
_SUPPORT_DEPARTMENTS = [d for d in DEPARTMENT_HEADCOUNT if d not in PRODUCTION_DEPARTMENTS]
_SUPPORT_HEADCOUNT_TOTAL = sum(DEPARTMENT_HEADCOUNT[d] for d in _SUPPORT_DEPARTMENTS)
DEPARTMENT_ALLOCATION_WEIGHT = {
    d: DEPARTMENT_HEADCOUNT[d] / _SUPPORT_HEADCOUNT_TOTAL for d in _SUPPORT_DEPARTMENTS
}

# Illustrative average daily salary cost per department (AED) — no real
# payroll data exists yet, so these are placeholders only. Feeds
# cost-per-productive-hour and revenue-per-labour-dirham in the Employee
# Portal's labour economics section. TODO: CONFIRM WITH GM/HR.
DEPARTMENT_AVG_DAILY_SALARY_AED = {
    "Baklava": 160.0,
    "French Bakery": 160.0,
    "Arabic Bread": 150.0,
    "Viennoiserie (Croissant/Danish)": 165.0,
    "Tahina": 140.0,
    "QC": 145.0,
    "Packing & Dispatch": 120.0,
    "Store": 120.0,
    "Maintenance": 150.0,
    "GM": 400.0,
    "Office": 180.0,
    "Drivers": 130.0,
    "Third-Party Cleaners": 100.0,
}

# Illustrative sample employees per department (placeholders — swap in
# Grandiose's real staff names/IDs once available). Not every headcount
# above has a named entry; these exist so the portal and department
# comparisons are demoable today.
EMPLOYEES = [
    {"name": "Ahmad Yousef", "id": "BAK-01", "department": "Baklava"},
    {"name": "Rania Haddad", "id": "BAK-02", "department": "Baklava"},
    {"name": "Mohammed Iqbal", "id": "BAK-03", "department": "Baklava"},

    {"name": "Rahul Nair", "id": "FRB-01", "department": "French Bakery"},
    {"name": "Fatima Zahra", "id": "FRB-02", "department": "French Bakery"},
    {"name": "John Cruz", "id": "FRB-03", "department": "French Bakery"},

    {"name": "Khalid Obaid", "id": "ARB-01", "department": "Arabic Bread"},
    {"name": "Youssef Amer", "id": "ARB-02", "department": "Arabic Bread"},
    {"name": "Layla Mansour", "id": "ARB-03", "department": "Arabic Bread"},

    {"name": "Ana Reyes", "id": "VIE-01", "department": "Viennoiserie (Croissant/Danish)"},
    {"name": "Omar Saleh", "id": "VIE-02", "department": "Viennoiserie (Croissant/Danish)"},
    {"name": "Priya Sharma", "id": "VIE-03", "department": "Viennoiserie (Croissant/Danish)"},

    {"name": "Hassan Ali", "id": "TAH-01", "department": "Tahina"},
    {"name": "Meera Pillai", "id": "TAH-02", "department": "Tahina"},

    {"name": "Noura Saeed", "id": "QC-01", "department": "QC"},
    {"name": "Vikram Singh", "id": "QC-02", "department": "QC"},

    {"name": "Josie Santos", "id": "PKD-01", "department": "Packing & Dispatch"},
    {"name": "Ravi Kumar", "id": "PKD-02", "department": "Packing & Dispatch"},
    {"name": "Abdullah Nasser", "id": "PKD-03", "department": "Packing & Dispatch"},

    {"name": "Aisha Rahman", "id": "STR-01", "department": "Store"},
    {"name": "Manuel Cruz", "id": "STR-02", "department": "Store"},

    {"name": "Suresh Pillai", "id": "MNT-01", "department": "Maintenance"},
    {"name": "Ibrahim Al Farsi", "id": "MNT-02", "department": "Maintenance"},

    {"name": "Khalifa Al Marzooqi", "id": "GM-01", "department": "GM"},

    {"name": "Sara Al Ali", "id": "OFF-01", "department": "Office"},
    {"name": "Deepak Verma", "id": "OFF-02", "department": "Office"},
    {"name": "Nadia Haddad", "id": "OFF-03", "department": "Office"},

    {"name": "Mohammed Rafiq", "id": "DRV-01", "department": "Drivers"},
    {"name": "Carlos Dionisio", "id": "DRV-02", "department": "Drivers"},
    {"name": "Salim Bakhit", "id": "DRV-03", "department": "Drivers"},

    {"name": "Ganesh Kumar", "id": "CLN-01", "department": "Third-Party Cleaners"},
    {"name": "Rosa Villanueva", "id": "CLN-02", "department": "Third-Party Cleaners"},
]


@st.cache_resource
def get_supabase_client():
    try:
        url = st.secrets["SUPABASE_URL"]
        key = st.secrets["SUPABASE_KEY"]
    except (KeyError, FileNotFoundError):
        return None
    try:
        return create_client(url, key)
    except Exception:
        return None

def portal_insert(table, row):
    client = get_supabase_client()
    if client is not None:
        try:
            client.table(table).insert(row).execute()
            return True, None
        except Exception as e:
            return False, str(e)
    else:
        row = {**row, "created_at": datetime.now().isoformat()}
        key = f"_local_{table}"
        st.session_state.setdefault(key, [])
        row["id"] = len(st.session_state[key]) + 1
        st.session_state[key].append(row)
        return True, None

def portal_fetch(table, filters=None):
    client = get_supabase_client()
    if client is not None:
        try:
            q = client.table(table).select("*")
            if filters:
                for k, v in filters.items():
                    q = q.eq(k, v)
            res = q.order("created_at", desc=True).execute()
            return pd.DataFrame(res.data), None
        except Exception as e:
            return pd.DataFrame(), str(e)
    else:
        rows = st.session_state.get(f"_local_{table}", [])
        df = pd.DataFrame(rows)
        if not df.empty:
            if filters:
                for k, v in filters.items():
                    df = df[df[k] == v]
            if "created_at" in df.columns:
                df = df.sort_values("created_at", ascending=False)
        return df, None

def portal_storage_mode():
    return "Supabase (persistent)" if get_supabase_client() is not None else "In-session only (not saved after restart)"


def compute_labour_summary(shift_df, department):
    """
    §A4-§A8 labour-efficiency chain for a group of shift-log rows (one
    employee's own logs, or a whole department's) — shared by the employee
    "My performance" tab and the Manager/HR department view so both always
    show the same numbers for the same underlying shifts.

    `shift_df` must already have its numeric columns coerced (the caller's
    `pd.to_numeric` pass). Returns a dict of the rolled-up hours (§A8:
    aggregated first, then converted to percentages) plus the labour
    economics figures — `None` for any ratio with a missing/zero input.
    """
    breaks_are_paid = LABOUR_CONFIG["breaks_are_paid"]
    shift_hours_rows = []
    for _, r in shift_df.iterrows():
        hw = r.get("hours_worked")
        paid_hours_row = hw if pd.notna(hw) and hw > 0 else LABOUR_CONFIG["shift_length_hours"]
        bm = r.get("break_minutes")
        bm = bm if pd.notna(bm) else LABOUR_CONFIG["default_break_minutes"]
        dm = r.get("downtime_minutes")
        dm = dm if pd.notna(dm) else 0
        cm = r.get("changeover_minutes")
        cm = cm if pd.notna(cm) else 0
        shift_hours_rows.append(labour_calc.compute_shift_hours(
            paid_hours=paid_hours_row, break_minutes=bm, changeover_minutes=cm,
            downtime_minutes=dm, breaks_are_paid=breaks_are_paid,
        ))
    rolled = labour_calc.rollup_hours(shift_hours_rows)

    naive_mean_true_efficiency = None
    per_row_true_eff = [labour_calc.true_efficiency_pct(sh["productive_hours"], sh["paid_hours"])
                         for sh in shift_hours_rows]
    per_row_true_eff = [v for v in per_row_true_eff if v is not None]
    if per_row_true_eff:
        naive_mean_true_efficiency = sum(per_row_true_eff) / len(per_row_true_eff)

    days_worked = len(shift_df)
    total_salary_cost = DEPARTMENT_AVG_DAILY_SALARY_AED.get(department)
    if total_salary_cost is not None:
        total_salary_cost = total_salary_cost * days_worked

    revenue_is_allocated = department not in PRODUCTION_DEPARTMENTS
    if not revenue_is_allocated:
        rev_series = shift_df["revenue_generated"] if "revenue_generated" in shift_df.columns else pd.Series(dtype=float)
        revenue_attributed = rev_series.sum() if rev_series.notna().any() else None
    else:
        all_perf, _err = portal_fetch("employee_performance")
        production_revenue_total = None
        if not all_perf.empty and {"revenue_generated", "department"} <= set(all_perf.columns):
            prod_rows = all_perf[all_perf["department"].isin(PRODUCTION_DEPARTMENTS)]
            prod_rev = pd.to_numeric(prod_rows["revenue_generated"], errors="coerce")
            production_revenue_total = prod_rev.sum() if prod_rev.notna().any() else None
        revenue_attributed = labour_calc.compute_revenue_attributed(
            production_revenue_total, DEPARTMENT_ALLOCATION_WEIGHT.get(department))

    return {
        "rolled": rolled,
        "naive_mean_true_efficiency": naive_mean_true_efficiency,
        "days_worked": days_worked,
        "breaks_are_paid": breaks_are_paid,
        "total_salary_cost": total_salary_cost,
        "revenue_is_allocated": revenue_is_allocated,
        "revenue_attributed": revenue_attributed,
        "cost_per_productive_hour": labour_calc.cost_per_productive_hour(total_salary_cost, rolled["productive_hours"]),
        "revenue_per_labour_dirham": labour_calc.revenue_per_labour_dirham(revenue_attributed, total_salary_cost),
        "revenue_per_day": labour_calc.revenue_per_day(revenue_attributed, days_worked),
    }


def render_labour_economics_cards(summary):
    """Shared 4-card labour-economics row (§A10.4) plus its muted footnotes —
    reused by both the employee and manager views so they render identically."""
    le1, le2, le3, le4 = st.columns(4)
    with le1:
        st.markdown(f"<div class='kpi-label'>Salary cost</div>"
                    f"<div class='kpi-value' style='font-size:1.4rem;'>{fmt_aed(summary['total_salary_cost'])}</div>",
                    unsafe_allow_html=True)
    with le2:
        st.markdown(f"<div class='kpi-label'>Cost / productive hour</div>"
                    f"<div class='kpi-value' style='font-size:1.4rem;'>{fmt_aed(summary['cost_per_productive_hour'])}</div>",
                    unsafe_allow_html=True)
    with le3:
        st.markdown(f"<div class='kpi-label'>Revenue attributed</div>"
                    f"<div class='kpi-value' style='font-size:1.4rem;'>{fmt_aed(summary['revenue_attributed'])}</div>",
                    unsafe_allow_html=True)
    with le4:
        st.markdown(f"<div class='kpi-label'>Revenue / labour dirham</div>"
                    f"<div class='kpi-value' style='font-size:1.4rem;'>{fmt_ratio(summary['revenue_per_labour_dirham'])}</div>",
                    unsafe_allow_html=True)
    if summary["revenue_is_allocated"]:
        st.caption("Revenue allocated, not directly attributed.")
    st.caption("Salary cost is an illustrative placeholder pending HR/finance-confirmed figures.")

    rev_per_day = summary["revenue_per_day"]
    if rev_per_day is None:
        st.caption(f"Actual AED/day: — vs GM benchmark "
                   f"AED {LABOUR_CONFIG['gm_daily_benchmark_aed']:,.0f}/day — variance: —")
    else:
        variance = rev_per_day - LABOUR_CONFIG["gm_daily_benchmark_aed"]
        variance_str = fmt_aed(variance)
        if variance >= 0:
            variance_str = f"+{variance_str}"
        st.caption(f"Actual AED/day: {fmt_aed(rev_per_day)} vs GM benchmark "
                   f"AED {LABOUR_CONFIG['gm_daily_benchmark_aed']:,.0f}/day — variance: {variance_str}")


def render_deduction_bar_and_table(rolled, breaks_are_paid):
    """
    §A10.2-3 — the proportional deduction bar and the deduction table below
    it. The bar's legend is custom HTML, not a Plotly legend: five segment
    names plus a "Hours" axis title and tick numbers all fighting for the
    same cramped strip is what made the Plotly legend unreadable, so the
    legend lives in the page below the chart instead, with the axis hidden
    entirely (the exact hours are already in the legend and the table).

    A segment with real but tiny hours (a few minutes of downtime) gets a
    minimum visual width so it isn't an invisible hairline — only the
    on-screen bar is padded, using hours borrowed from the largest segment;
    the legend and the deduction table always show the true, unpadded hours.
    """
    raw_segments = [
        ("Productive", rolled["bar_productive_hours"], COLORS["primary"]),
        ("Changeover", rolled["changeover_hours"], COLORS["primary_deep"]),
        ("Downtime", rolled["downtime_hours"], COLORS["warning"]),
        ("Breaks", rolled["break_hours"], COLORS["secondary"]),
        ("Idle", rolled["idle_waiting_hours"], COLORS["tertiary"]),
    ]
    total_hours = sum(hrs for _, hrs, _ in raw_segments) or 1.0
    floor_hours = 0.035 * total_hours
    display_hours = []
    pad_total = 0.0
    for _, hrs, _ in raw_segments:
        if 0 < hrs < floor_hours:
            display_hours.append(floor_hours)
            pad_total += floor_hours - hrs
        else:
            display_hours.append(hrs)
    if pad_total > 0:
        largest_idx = max(range(len(raw_segments)), key=lambda i: raw_segments[i][1])
        display_hours[largest_idx] = max(display_hours[largest_idx] - pad_total, 0.0)

    fig_bar = go.Figure()
    for (label, hrs, color), disp in zip(raw_segments, display_hours):
        fig_bar.add_trace(go.Bar(
            x=[disp], y=[""], orientation="h", marker_color=color, marker_line_width=0,
            showlegend=False, hovertext=[f"{label}: {fmt_hours(hrs)}h"], hoverinfo="text",
        ))
    fig_bar.update_layout(**PLOTLY_DARK, height=54, barmode="stack", showlegend=False,
                          margin=dict(l=4, r=4, t=4, b=4),
                          xaxis=dict(visible=False), yaxis=dict(visible=False))
    st.plotly_chart(fig_bar, width='stretch', config={"displayModeBar": False})

    legend_html = "".join(
        f"<span style='display:inline-flex; align-items:center; gap:5px; margin:2px 16px 2px 0;'>"
        f"<span style='width:10px; height:10px; border-radius:2px; background-color:{color}; "
        f"display:inline-block; flex-shrink:0;'></span>"
        f"<span style='color:{COLORS['text_soft']}; font-size:0.76rem;'>{label} ({fmt_hours(hrs)}h)</span></span>"
        for label, hrs, color in raw_segments
    )
    st.markdown(f"<div style='display:flex; flex-wrap:wrap; margin:2px 0 14px 0;'>{legend_html}</div>",
                unsafe_allow_html=True)

    ded_rows = [
        ("Paid hours", rolled["paid_hours"], False),
        ("Break time (paid, not deducted)" if breaks_are_paid else "− Break time",
         rolled["break_hours"], False),
        ("− Changeover / setup", rolled["changeover_hours"], False),
        ("− Waiting for equipment", rolled["downtime_hours"], False),
        ("− Idle waiting", rolled["idle_waiting_hours"], False),
        ("= Productive hours", rolled["productive_hours"], True),
    ]
    ded_html = "".join(
        f"<tr style='border-top:{('1px solid ' + COLORS['border']) if bordered else 'none'};"
        f"font-weight:{700 if bordered else 400};'>"
        f"<td style='padding:4px 12px 4px 0;'>{label}</td>"
        f"<td style='padding:4px 0; text-align:right; "
        f"font-family:ui-monospace,SFMono-Regular,Menlo,monospace;'>{fmt_hours(hrs)}</td></tr>"
        for label, hrs, bordered in ded_rows
    )
    st.markdown(f"<table style='width:100%; border-collapse:collapse; color:{COLORS['text']};'>"
                f"{ded_html}</table>", unsafe_allow_html=True)


# ----------------------------------------------------------------------
# APP CHROME — top navigation bar, page header, footer
#
# Navigation is a single horizontal bar (logo | section links | icons),
# matching the design system's nav spec: uppercase label-sm links, muted
# by default, gold with a 2px gold underline when active. Each entry in
# NAV_ITEMS is (key, label); `key` is what the section routing below
# compares against, so labels can be reworded without touching routing.
# ----------------------------------------------------------------------
NAV_ITEMS = [
    ("performance_tracker", "Performance Tracker"),
    ("scenario_resilience", "Scenario & Resilience"),
    ("employee_portal",     "Employee Portal"),
    ("company_profile",     "Company Profile"),
    ("sku_performance",     "SKU Performance"),
    ("data_processor",      "Data Processor"),
]

# Per-section page header: (title, subtitle). Rendered directly under the
# nav bar so every section opens with the same masthead treatment.
PAGE_META = {
    "performance_tracker": ("Financial Performance Dashboard",
                            "Where the bakery division stands today."),
    "scenario_resilience": ("Scenario & Resilience",
                            "What happens if — stress-tests the same cost and margin outputs "
                            "under new assumptions."),
    "employee_portal":     ("Employee Portal",
                            "Employee self-service performance tracking — bakery division."),
    "company_profile":     ("Company Profile",
                            "Everything collected from the Grandiose team so far — from the "
                            "27 Jul personal meeting with the GM."),
    "sku_performance":     ("Bakery Product Performance",
                            "Product and SKU performance across the six bakery divisions."),
    "data_processor":      ("Data Processor",
                            "Drop in rough Excel, PDF, or Word files — get back a cleaned, "
                            "evaluated Excel report."),
}

FOOTER_LINKS = ["Privacy Policy", "Terms of Service", "Financial Disclosure", "Contact Support"]

# Shared height for every cell in the nav row — wordmark, links and icons all
# get this exact box so they resolve to one common centre line regardless of
# how tall their own content happens to be.
NAV_ROW_H = 38

NAV_ICONS_HTML = (
    "<div class='gd-nav-icons'>"
    "<svg viewBox='0 0 24 24' role='img' aria-label='Notifications'>"
    "<path d='M12 22c1.1 0 2-.9 2-2h-4c0 1.1.9 2 2 2zm6-6v-5c0-3.07-1.63-5.64-4.5-6.32V4c0-.83-.67-1.5-1.5-1.5"
    "s-1.5.67-1.5 1.5v.68C7.64 5.36 6 7.92 6 11v5l-2 2v1h16v-1l-2-2z'/></svg>"
    "<svg viewBox='0 0 24 24' role='img' aria-label='Account'>"
    "<path d='M12 2C6.48 2 2 6.48 2 12s4.48 10 10 10 10-4.48 10-10S17.52 2 12 2zm0 3c1.66 0 3 1.34 3 3s-1.34 3-3 3"
    "-3-1.34-3-3 1.34-3 3-3zm0 14.2c-2.5 0-4.71-1.28-6-3.22.03-1.99 4-3.08 6-3.08 1.99 0 5.97 1.09 6 3.08"
    "-1.29 1.94-3.5 3.22-6 3.22z'/></svg>"
    "</div>"
)

st.markdown(f"""
<style>
    .stApp {{
        background-image: radial-gradient(1200px 420px at 15% -10%, {hex_to_rgba(COLORS['primary'], 0.08)} 0%, transparent 60%);
        background-attachment: fixed;
    }}
    .gold-divider {{ height: 2px; width: 32px; background-color: {COLORS['primary_deep']}; margin-bottom: 14px; }}

    /* Pull the canvas up under the nav bar and give the page real gutters,
       so the bar reads as full-bleed app chrome rather than page content. */
    .stMainBlockContainer {{ padding-top: 2.2rem !important; max-width: 1350px; }}

    /* ---------------- TOP NAVIGATION BAR ---------------- */
    /* A single bar with everything on one centre line — logo, links and
       icons all sit on the same axis, the way the design has it. */
    .st-key-gd_topbar {{
        position: sticky; top: 0; z-index: 999;
        background-color: {COLORS['bg']};
        border-bottom: 1px solid {COLORS['border']};
        margin-bottom: 34px;
        padding: 12px 0 10px 0;
    }}
    /* Streamlit wraps each markdown block in containers that carry their own
       margins. Left alone they make the wordmark/icon cells taller than the
       button cells, so centring resolves to a different line for each and the
       row looks off. Zero them out inside the bar only. */
    .st-key-gd_topbar [data-testid="stMarkdown"],
    .st-key-gd_topbar [data-testid="stMarkdownContainer"],
    .st-key-gd_topbar [data-testid="stElementContainer"],
    .st-key-gd_topbar [data-testid="stVerticalBlock"] {{
        margin: 0 !important; padding: 0 !important; gap: 0 !important;
    }}
    /* Nav links: strip the global ghost-button chrome down to a text link. */
    .st-key-gd_topbar [data-testid^="stBaseButton"] {{
        background-color: transparent !important;
        border: none !important;
        border-bottom: 2px solid transparent !important;
        border-radius: 0 !important;
        /* Symmetric padding + a fixed height so the link text lands on the
           same centre line as the wordmark and the icons. */
        padding: 0 !important;
        height: {NAV_ROW_H}px !important; min-height: 0 !important;
        display: flex !important; align-items: center !important;
        justify-content: center !important;
        transition: box-shadow 0.25s ease, color 0.25s ease;
    }}
    .st-key-gd_topbar [data-testid^="stBaseButton"] p {{
        color: {COLORS['text_soft']} !important;
        font-size: 0.63rem !important;
        font-weight: 600 !important;
        letter-spacing: 0.10em !important;
        text-transform: uppercase !important;
        white-space: nowrap;
    }}
    /* Underline is drawn as an inset shadow rather than a border, so the
       active state cannot add 2px of height and knock the row off centre. */
    .st-key-gd_topbar [data-testid^="stBaseButton"]:hover {{
        background-color: transparent !important;
        box-shadow: inset 0 -2px 0 0 {hex_to_rgba(COLORS['primary'], 0.45)} !important;
    }}
    .st-key-gd_topbar [data-testid^="stBaseButton"]:hover p {{ color: {COLORS['primary']} !important; }}
    /* Active link — outranks the global solid-gold primary-button rule
       because this selector adds the container class. */
    .st-key-gd_topbar [data-testid*="rimary" i] {{
        background-color: transparent !important;
        box-shadow: inset 0 -2px 0 0 {COLORS['primary']} !important;
    }}
    .st-key-gd_topbar [data-testid*="rimary" i]:hover {{ background-color: transparent !important; }}
    .st-key-gd_topbar [data-testid*="rimary" i] p {{
        color: {COLORS['primary']} !important; font-weight: 700 !important;
    }}

    /* Wordmark sits hard against the left edge of the canvas — the negative
       pull cancels the column's own gutter so it lines up with the very
       start of the bar rather than floating inboard of it. */
    .gd-logo {{
        font-family: 'Bebas Neue', sans-serif !important;
        font-size: 2.45rem; line-height: 1; color: {COLORS['primary']};
        text-transform: uppercase; letter-spacing: 0.13em;
        padding: 0 16px 0 0;
        white-space: nowrap;
        display: flex; align-items: center; height: {NAV_ROW_H}px;
    }}
    /* Status icons — inline SVG rather than an icon font, so they can never
       flash their ligature text ("notifications") while a font loads. */
    .gd-nav-icons {{
        display: flex; justify-content: flex-end; align-items: center;
        gap: 18px; padding: 0; height: {NAV_ROW_H}px;
    }}
    .gd-nav-icons svg {{
        width: 18px; height: 18px; flex: none;
        fill: {COLORS['text_soft']}; transition: fill 0.2s ease;
    }}
    .gd-nav-icons svg:hover {{ fill: {COLORS['primary']}; }}

    /* ---------------- PAGE HEADER ---------------- */
    .stApp .gd-page-title, .stApp .gd-page-title span {{
        font-family: 'Bebas Neue', sans-serif !important;
        font-size: 3.1rem; line-height: 1.02; color: {COLORS['text']} !important;
        text-transform: uppercase; letter-spacing: 0.01em; margin: 0;
    }}
    .gd-page-sub {{
        color: {COLORS['text_soft']}; font-size: 0.88rem; margin-top: 8px;
    }}
    .gd-page-header {{
        border-bottom: 1px solid {COLORS['border']};
        padding-bottom: 20px; margin-bottom: 34px;
    }}

    /* ---------------- FOOTER ---------------- */
    .gd-footer {{
        border-top: 1px solid {COLORS['border']};
        background-color: {COLORS['surface_lowest']};
        margin-top: 56px; padding: 26px 28px;
        display: flex; flex-wrap: wrap; gap: 14px;
        justify-content: space-between; align-items: center;
    }}
    .gd-footer-copy {{
        color: {COLORS['text_soft']} !important;
        font-size: 0.68rem; font-weight: 600;
        text-transform: uppercase; letter-spacing: 0.1em;
    }}
    .gd-footer-links {{ display: flex; flex-wrap: wrap; gap: 22px; }}
    .gd-footer-links span {{
        color: {COLORS['text_soft']} !important;
        font-size: 0.68rem; font-weight: 600;
        text-transform: uppercase; letter-spacing: 0.1em;
        cursor: default; transition: color 0.2s ease;
    }}
    .gd-footer-links span:hover {{ color: {COLORS['primary']} !important; }}
    .gd-footer-note {{
        color: {COLORS['text_soft']} !important;
        font-size: 0.72rem; line-height: 1.6;
        padding: 16px 28px 0 28px; opacity: 0.75;
    }}
</style>
""", unsafe_allow_html=True)


def render_top_nav():
    """Full-width top bar: logo, section links, status icons.

    Returns the key of the active section. Selection is held in
    st.session_state so it survives the reruns that widgets trigger.
    """
    st.session_state.setdefault("active_section", NAV_ITEMS[0][0])

    # Link columns are sized from their label length so adding or renaming a
    # section rebalances the bar automatically instead of overflowing it.
    link_widths = [0.30 + 0.088 * len(label) for _, label in NAV_ITEMS]

    with st.container(key="gd_topbar"):
        cols = st.columns([2.1, *link_widths, 0.7], vertical_alignment="center")
        cols[0].markdown("<div class='gd-logo'>Grandiose</div>", unsafe_allow_html=True)

        for col, (key, label) in zip(cols[1:], NAV_ITEMS):
            is_active = st.session_state["active_section"] == key
            if col.button(label, key=f"nav_{key}", width='stretch',
                          type="primary" if is_active else "secondary"):
                st.session_state["active_section"] = key
                st.rerun()

        cols[-1].markdown(NAV_ICONS_HTML, unsafe_allow_html=True)

    return st.session_state["active_section"]


def render_page_header(section_key):
    title, subtitle = PAGE_META[section_key]
    st.markdown(
        f"<div class='gd-page-header'>"
        f"<div class='gold-divider'></div>"
        f"<h1 class='gd-page-title'>{title}</h1>"
        f"<div class='gd-page-sub'>{subtitle}</div>"
        f"</div>",
        unsafe_allow_html=True,
    )


def render_footer():
    links = "".join(f"<span>{name}</span>" for name in FOOTER_LINKS)
    st.markdown(
        f"<div class='gd-footer'>"
        f"<span class='gd-footer-copy'>© 2024 Grandiose. Institutional rigor. Culinary artistry.</span>"
        f"<div class='gd-footer-links'>{links}</div>"
        f"</div>"
        f"<div class='gd-footer-note'>"
        f"Financial performance and cost optimization for Grandiose Bakery operations · GIP III · "
        f"Bakery division only, catering excluded · Figures shown are illustrative benchmarks "
        f"pending Grandiose-provided actuals."
        f"</div>",
        unsafe_allow_html=True,
    )


# ----------------------------------------------------------------------
# BAKERY PRODUCT PERFORMANCE — helpers
#
# The page renders one bubble per SKU, grouped into a cluster per division.
# Bubble AREA is proportional to units sold (so radius scales with the
# square root), which is what makes relative volume readable at a glance.
# ----------------------------------------------------------------------

# One gold ramp rather than six arbitrary hues — divisions stay separable
# without introducing colours the design system doesn't have.
SKU_DIVISION_COLORS = {
    "Baklava":             "#F2CA50",
    "French Bakery":       "#D4AF37",
    "Arabic Bread":        "#B8942C",
    "Viennoiserie":        "#9C7A22",
    "Tahina":              "#7E621B",
    "Seasonal Collection": "#FFE088",
}

BUBBLE_W, BUBBLE_H = 1200, 430
BUBBLE_R_MIN, BUBBLE_R_MAX = 6.5, 36.0


@st.cache_data(show_spinner=False)
def load_sku_products():
    return sku_data.load_products()


def _bubble_radius(units, max_units):
    """Radius for a bubble, scaled so AREA tracks units sold."""
    if max_units <= 0:
        return BUBBLE_R_MIN
    return BUBBLE_R_MIN + (BUBBLE_R_MAX - BUBBLE_R_MIN) * math.sqrt(units / max_units)


@st.cache_data(show_spinner=False)
def bubble_layout(items, width=BUBBLE_W, height=BUBBLE_H, pad=4.0):
    """Pack bubbles into one organic cluster per division.

    `items` is a tuple of (cluster_index, radius) — a tuple so the result
    caches, which also keeps positions stable across reruns instead of
    reshuffling every time the page redraws.

    Seeds each cluster on a phyllotaxis spiral (largest bubbles innermost,
    which reads as organic rather than gridded), then relaxes overlaps while
    a weak spring holds each bubble to its own cluster so the groups stay
    visually distinct.
    """
    if not items:
        return []

    members = {}
    for i, (cluster, _r) in enumerate(items):
        members.setdefault(cluster, []).append(i)

    clusters = sorted(members)
    margin = BUBBLE_R_MAX + 16
    usable = max(width - 2 * margin, 1)
    step = usable / len(clusters)
    centres = {c: (margin + (k + 0.5) * step, height / 2.0)
               for k, c in enumerate(clusters)}

    golden = math.pi * (3.0 - math.sqrt(5.0))
    pos = [[0.0, 0.0] for _ in items]

    for cluster in clusters:
        idxs = sorted(members[cluster], key=lambda i: -items[i][1])
        cx, cy = centres[cluster]
        area = sum(math.pi * items[i][1] ** 2 for i in idxs)
        # Capped against the per-cluster slot so neighbouring divisions keep
        # a visible gap instead of merging into one mass.
        spread = min(math.sqrt(area / math.pi) * 1.15, step * 0.42)
        for j, i in enumerate(idxs):
            angle = j * golden
            radius = spread * math.sqrt((j + 0.35) / len(idxs))
            pos[i] = [cx + radius * math.cos(angle), cy + radius * math.sin(angle)]

    n = len(items)
    for _ in range(160):
        for i in range(n):
            for j in range(i + 1, n):
                dx = pos[j][0] - pos[i][0]
                dy = pos[j][1] - pos[i][1]
                dist = math.hypot(dx, dy)
                min_dist = items[i][1] + items[j][1] + pad
                if dist < min_dist:
                    if dist < 1e-6:
                        dx, dy, dist = 0.7, 0.7, 1.0
                    shift = (min_dist - dist) / dist * 0.5
                    ox, oy = dx * shift, dy * shift
                    pos[i][0] -= ox; pos[i][1] -= oy
                    pos[j][0] += ox; pos[j][1] += oy

        for i, (cluster, r) in enumerate(items):
            cx, cy = centres[cluster]
            pos[i][0] += (cx - pos[i][0]) * 0.022
            pos[i][1] += (cy - pos[i][1]) * 0.030
            pos[i][0] = min(max(pos[i][0], r + 2), width - r - 2)
            pos[i][1] = min(max(pos[i][1], r + 2), height - r - 2)

    return [(round(x, 2), round(y, 2)) for x, y in pos]


def build_bubble_figure(products, selected_sku=None):
    """Bubble clusters, one Scatter trace per division."""
    max_units = max((p["units"] for p in products), default=1)
    divisions = [d for d in sku_data.DIVISIONS
                 if any(p["division"] == d for p in products)]
    div_index = {d: i for i, d in enumerate(divisions)}

    ordered = sorted(products, key=lambda p: (div_index[p["division"]], -p["units"]))
    radii = [_bubble_radius(p["units"], max_units) for p in ordered]
    coords = bubble_layout(tuple((div_index[p["division"]], r)
                                 for p, r in zip(ordered, radii)))

    fig = go.Figure()
    for division in divisions:
        rows = [(p, r, xy) for p, r, xy in zip(ordered, radii, coords)
                if p["division"] == division]
        if not rows:
            continue
        is_seasonal = division == sku_data.SEASONAL_DIVISION

        custom = [[
            p["product"], p["division"], f"{p['units']:,}",
            sku_data.format_aed(p["sales_aed"]), f"{p['contribution_pct']:.1f}%",
            p["rank"], p["sku"],
            p.get("collection", ""), p.get("availability", ""),
        ] for p, _r, _xy in rows]

        detail = (
            "<b>%{customdata[0]}</b><br>"
            "<span style='color:#F2CA50'>%{customdata[1]}</span><br><br>"
            "Units sold&nbsp;&nbsp;<b>%{customdata[2]}</b><br>"
            "Sales value&nbsp;&nbsp;<b>%{customdata[3]}</b><br>"
            "Contribution&nbsp;&nbsp;<b>%{customdata[4]}</b><br>"
            "Rank&nbsp;&nbsp;<b>#%{customdata[5]}</b>"
        )
        if is_seasonal:
            detail += ("<br><br>Collection&nbsp;&nbsp;<b>%{customdata[7]}</b>"
                       "<br>Available&nbsp;&nbsp;<b>%{customdata[8]}</b>")

        fig.add_trace(go.Scatter(
            x=[xy[0] for _p, _r, xy in rows],
            y=[xy[1] for _p, _r, xy in rows],
            mode="markers", name=division, customdata=custom,
            marker=dict(
                size=[r * 2 for _p, r, _xy in rows], sizemode="diameter",
                color=SKU_DIVISION_COLORS.get(division, COLORS["primary"]),
                opacity=0.9,
                line=dict(width=1, color=hex_to_rgba(COLORS["bg"], 0.85)),
            ),
            hovertemplate=detail + "<extra></extra>",
        ))

    # Ring the selected SKU rather than restyling its marker in place.
    if selected_sku:
        for p, r, xy in zip(ordered, radii, coords):
            if p["sku"] == selected_sku:
                fig.add_trace(go.Scatter(
                    x=[xy[0]], y=[xy[1]], mode="markers", showlegend=False,
                    hoverinfo="skip",
                    marker=dict(size=r * 2 + 12, sizemode="diameter",
                                color="rgba(0,0,0,0)",
                                line=dict(width=2, color=COLORS["text"])),
                ))
                break

    for division in divisions:
        rows = [xy for p, _r, xy in zip(ordered, radii, coords)
                if p["division"] == division]
        if rows:
            fig.add_annotation(
                x=sum(x for x, _y in rows) / len(rows), y=-14,
                text=division.upper(), showarrow=False,
                font=dict(family="Inter", size=10, color=COLORS["text_soft"]),
            )

    fig.update_layout(
        height=470, showlegend=False,
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        margin=dict(l=8, r=8, t=10, b=34),
        xaxis=dict(visible=False, range=[0, BUBBLE_W], fixedrange=True),
        yaxis=dict(visible=False, range=[-30, BUBBLE_H], fixedrange=True),
        hoverlabel=dict(bgcolor=COLORS["surface"], bordercolor=COLORS["border"],
                        font=dict(family="Inter", size=12, color=COLORS["text"]),
                        align="left"),
        dragmode=False,
    )
    return fig


def sku_table(products, division):
    """Display-ready SKU table for one division."""
    rows = sorted((p for p in products if p["division"] == division),
                  key=lambda p: p["rank"])
    table = {
        "SKU": [p["sku"] for p in rows],
        "Product": [p["product"] for p in rows],
        "Units sold": [f"{p['units']:,}" for p in rows],
        "Sales value": [sku_data.format_aed(p["sales_aed"]) for p in rows],
        "Contr. %": [f"{p['contribution_pct']:.1f}%" for p in rows],
        "Rank": [p["rank"] for p in rows],
    }
    if division == sku_data.SEASONAL_DIVISION:
        table["Availability"] = [p.get("availability", "—") for p in rows]
    return pd.DataFrame(table)


section = render_top_nav()
render_page_header(section)

# ========================================================================
# SECTION A — PERFORMANCE TRACKER
# ========================================================================
if section == "performance_tracker":
    kpis = [
        ("🍞", "Food cost %", f"{baseline['food_cost_pct']}%", "↑ 1.2pt vs target", "pill-up-bad",
         food_cost_trend, COLORS["primary"]),
        ("🗑️", "Wastage %", f"{baseline['wastage_pct']}%", "↓ 0.3pt vs last month", "pill-down-good",
         wastage_trend, COLORS["success"]),
        ("📈", "Gross margin", f"{baseline['gross_margin_pct']}%", "flat vs last month", "pill-flat",
         margin_trend, COLORS["tertiary"]),
        ("💵", "Cost / unit (AED)", f"{baseline['cost_per_unit']:.2f}", f"↑ vs std {baseline['standard_cost_per_unit']:.2f}", "pill-up-bad",
         cost_unit_trend, COLORS["secondary"]),
    ]
    cols = st.columns(4)
    for col, (icon, label, value, delta, pill_class, trend, color) in zip(cols, kpis):
        with col:
            with st.container(border=True):
                st.markdown(f"""
                <div style='padding:10px 14px 0 14px;'>
                    <span class='kpi-icon' style='background-color:{hex_to_rgba(color,0.15)};'>{icon}</span>
                    <span class='kpi-label'>{label}</span>
                    <div class='kpi-value'>{value}</div>
                    <span class='pill {pill_class}'>{delta}</span>
                </div>
                """, unsafe_allow_html=True)
                st.plotly_chart(sparkline(trend, color), width='stretch',
                                 config={"displayModeBar": False}, key=f"spark_{label}")

    st.markdown("#### Efficiency at a glance")
    colG1, colG2 = st.columns(2)
    with colG1:
        with st.container(border=True):
            st.markdown(f"<div class='kpi-label'>Wastage vs. 1% target</div>", unsafe_allow_html=True)
            fig_waste_gauge = go.Figure(go.Indicator(
                mode="gauge+number",
                value=baseline['wastage_pct'],
                number={'suffix': '%', 'font': {'color': COLORS['text'], 'size': 30, 'family': 'Bebas Neue'}},
                gauge={
                    'axis': {'range': [0, 8], 'tickcolor': COLORS['text_soft'],
                             'tickfont': {'color': COLORS['text_soft'], 'size': 9}, 'tickwidth': 1},
                    'bar': {'color': COLORS['primary'], 'thickness': 0.28},
                    'bgcolor': 'rgba(0,0,0,0)',
                    'borderwidth': 0,
                    'steps': [
                        {'range': [0, WASTAGE_TARGET_PCT], 'color': hex_to_rgba(COLORS['success'], 0.22)},
                        {'range': [WASTAGE_TARGET_PCT, 3], 'color': hex_to_rgba(COLORS['warning'], 0.22)},
                        {'range': [3, 8], 'color': hex_to_rgba(COLORS['danger'], 0.22)},
                    ],
                    'threshold': {'line': {'color': COLORS['primary_deep'], 'width': 3},
                                  'thickness': 0.85, 'value': WASTAGE_TARGET_PCT},
                },
            ))
            fig_waste_gauge.update_layout(**PLOTLY_DARK, height=210)
            st.plotly_chart(fig_waste_gauge, width='stretch', config={"displayModeBar": False})
            st.caption("Gold needle line marks the 1% company target · current is above target.")

    with colG2:
        with st.container(border=True):
            st.markdown(f"<div class='kpi-label'>Cost structure (% of revenue)</div>", unsafe_allow_html=True)
            cost_labels = ["Food cost", "Labour (target)", "Packaging", "Overhead", "Margin & other"]
            cost_values = [baseline['food_cost_pct'], 20.0, 2.1, 9.4,
                            max(100 - baseline['food_cost_pct'] - 20.0 - 2.1 - 9.4, 0)]
            fig_cost_donut = go.Figure(go.Pie(
                labels=cost_labels, values=cost_values, hole=0.62,
                marker=dict(colors=CATEGORICAL, line=dict(color=COLORS["bg"], width=2)),
                textfont=dict(color=COLORS["text"], size=10),
            ))
            fig_cost_donut.update_layout(
                **PLOTLY_DARK, height=210, showlegend=True,
                legend=dict(orientation="h", y=-0.15, font=dict(color=COLORS["text_soft"], size=9)),
                annotations=[dict(text=f"<b>{baseline['gross_margin_pct']}%</b><br><span style='font-size:9px;color:{COLORS['text_soft']}'>margin</span>",
                                   x=0.5, y=0.5, font=dict(size=16, color=COLORS['primary']), showarrow=False)],
            )
            st.plotly_chart(fig_cost_donut, width='stretch', config={"displayModeBar": False})
            st.caption("Illustrative allocation, blending current actuals with target labour cost.")

    st.markdown("#### Food cost % vs target — last 6 months")
    trend_view = st.radio("View", ["Historical", "Projected (+2 months)"], horizontal=True,
                           label_visibility="collapsed", key="trend_view_toggle")
    with st.container(border=True):
        fig = go.Figure()
        plot_months, plot_trend = list(months), list(food_cost_trend)
        plot_target = list(target_line)
        is_projected = trend_view.startswith("Projected")
        if is_projected:
            slope = food_cost_trend[-1] - food_cost_trend[-2]
            proj_vals = [round(food_cost_trend[-1] + slope * i, 1) for i in (1, 2)]
            plot_months = plot_months + ["Aug*", "Sep*"]
            plot_trend = plot_trend + proj_vals
            plot_target = plot_target + [target_line[-1], target_line[-1]]
        # glow effect: wide low-opacity trace behind crisp line
        fig.add_trace(go.Scatter(x=plot_months, y=plot_trend, mode="lines", line=dict(color=COLORS["primary"], width=10),
                                  opacity=0.15, showlegend=False, hoverinfo="skip"))
        line_dash = "dot" if is_projected else "solid"
        fig.add_trace(go.Scatter(x=plot_months[:6], y=plot_trend[:6], mode="lines+markers", name="Food cost %",
                                  line=dict(color=COLORS["primary"], width=3),
                                  marker=dict(size=8, color=COLORS["bg"], line=dict(color=COLORS["primary"], width=2)),
                                  fill="tozeroy", fillcolor=hex_to_rgba(COLORS["primary"], 0.12)))
        if is_projected:
            fig.add_trace(go.Scatter(x=plot_months[5:], y=plot_trend[5:], mode="lines+markers", name="Projected",
                                      line=dict(color=COLORS["primary_deep"], width=3, dash="dot"),
                                      marker=dict(size=7, symbol="diamond", color=COLORS["bg"],
                                                  line=dict(color=COLORS["primary_deep"], width=2))))
        fig.add_trace(go.Scatter(x=plot_months, y=plot_target, mode="lines", name="Target",
                                  line=dict(color=COLORS["text_soft"], width=1.5, dash="dot")))
        fig.update_layout(**PLOTLY_DARK, height=300,
                           yaxis=dict(ticksuffix="%", range=[28, 34], gridcolor=GRID_COLOR, zeroline=False),
                           xaxis=dict(gridcolor="rgba(0,0,0,0)"),
                           legend=dict(orientation="h", yanchor="bottom", y=1.03, x=0,
                                       font=dict(color=COLORS["text_soft"])))
        st.plotly_chart(fig, width='stretch', config={"displayModeBar": False})
        if is_projected:
            st.caption("* Projected months extrapolate the current trend line — illustrative, not a forecast model.")
    render_email_share("trend_chart", "Email this Performance Tracker report", section="performance_tracker")

    st.markdown("#### Category panels")
    with st.container(key="gd_catpanels"):
        cols2 = st.columns(4)
        for col, (name, metrics) in zip(cols2, category_panels.items()):
            with col:
                with st.container(border=True):
                    rows = "".join(
                        f"<div class='cat-metric'>"
                        f"<div class='cat-metric-label'>{k}</div>"
                        f"<div class='cat-metric-value'>{v}</div>"
                        f"</div>"
                        for k, v in metrics.items()
                    )
                    st.markdown(f"<div class='cat-panel'><h4>{name}</h4>{rows}</div>",
                                unsafe_allow_html=True)
    render_email_share("category_panels", "Email this Performance Tracker report", section="performance_tracker")

# ========================================================================
# TAB B — SCENARIO & RESILIENCE
# ========================================================================
# ========================================================================
# SECTION B — SCENARIO & RESILIENCE
# ========================================================================
elif section == "scenario_resilience":
    mod1, mod2, mod3, mod4 = st.tabs([
        "1 · Inflation sensitivity", "2 · Supply disruption risk",
        "3 · Pandemic preparedness", "4 · Supplier alternatives",
    ])

    # ---------------- MODULE 1: INFLATION-ADJUSTED COST SENSITIVITY ----------------
    with mod1:
        st.markdown("##### 🔥 Inflation-adjusted cost sensitivity")
        st.caption("Dual-track inflation input with an adjustable bar, net of any subsidy/offset deducted.")

        base_cost = baseline["cost_per_unit"]
        colL, colR = st.columns([1, 1.3])
        with colL:
            with st.container(border=True):
                headline_inf = st.slider("Headline inflation forecast (%)", 0.0, 10.0, 2.8, 0.1)
                food_inf = st.slider("Actual food-input inflation (%)", 0.0, 20.0, 5.3, 0.1,
                                      help="Flour, dairy, and packaging inflation can diverge sharply from headline CPI during shocks.")
                subsidy_offset = st.number_input("Subsidy / price-cap offset (AED per unit)", 0.0, 2.0, 0.05, 0.01)

                st.markdown(f"<div style='margin-top:6px; color:{COLORS['text_soft']}; font-size:0.82rem;'>"
                            f"Cost composition — updates live with the sliders above</div>", unsafe_allow_html=True)
                inflation_addon = round(base_cost * food_inf / 100, 3)
                net_after_subsidy = round(base_cost + inflation_addon - subsidy_offset, 2)
                donut = go.Figure(go.Pie(
                    labels=["Base cost", "Inflation add-on"],
                    values=[base_cost, max(inflation_addon, 0.001)],
                    hole=0.64, sort=False, textinfo="percent",
                    marker=dict(colors=[COLORS["text_soft"], COLORS["danger"]],
                                line=dict(color=COLORS["surface"], width=3)),
                ))
                donut.update_layout(
                    **PLOTLY_DARK, height=220,
                    legend=dict(orientation="h", y=-0.12, font=dict(color=COLORS["text_soft"])),
                    annotations=[dict(
                        text=f"<b>AED {net_after_subsidy:.2f}</b><br><span style='font-size:10px;color:{COLORS['text_soft']}'>net / unit</span>",
                        x=0.5, y=0.5, font=dict(size=18, color=COLORS['text']), showarrow=False)],
                )
                st.plotly_chart(donut, width='stretch', config={"displayModeBar": False}, key="inflation_donut")

        adj_cost_headline = round(base_cost * (1 + headline_inf / 100) - subsidy_offset, 3)
        adj_cost_food = round(base_cost * (1 + food_inf / 100) - subsidy_offset, 3)
        base_revenue_per_unit = base_cost / (baseline["food_cost_pct"] / 100)
        food_cost_pct_adj = round((adj_cost_food / base_revenue_per_unit) * 100, 1)

        with colR:
            m1, m2, m3 = st.columns(3)
            m1.metric("Cost/unit — headline scenario", f"AED {adj_cost_headline:.2f}")
            m2.metric("Cost/unit — food-inflation scenario", f"AED {adj_cost_food:.2f}",
                      f"{adj_cost_food - base_cost:+.2f} vs current")
            m3.metric("Food cost % — food-inflation scenario", f"{food_cost_pct_adj}%",
                      f"{food_cost_pct_adj - baseline['food_cost_pct']:+.1f}pt", delta_color="inverse")

            with st.container(border=True):
                fig1 = go.Figure(go.Bar(
                    x=["Current", "Headline scenario", "Food-inflation scenario"],
                    y=[base_cost, adj_cost_headline, adj_cost_food],
                    marker_color=[COLORS["text_soft"], COLORS["tertiary"], COLORS["danger"]],
                    marker_line_width=0,
                ))
                fig1.update_layout(**PLOTLY_DARK, height=260, yaxis_title="AED per unit",
                                    yaxis=dict(gridcolor=GRID_COLOR), xaxis=dict(gridcolor="rgba(0,0,0,0)"))
                st.plotly_chart(fig1, width='stretch', config={"displayModeBar": False})

        st.caption("Context: UAE food inflation has ranged from a 2022 peak near 9% to negative readings in 2021, "
                   "against headline forecasts around 2–3% for 2026 — a single flat assumption misses this swing.")

    # ---------------- MODULE 2: NATURAL CALAMITY / SUPPLY DISRUPTION ----------------
    with mod2:
        st.markdown("##### 🌪️ Natural calamity & supply disruption risk")
        st.caption("UAE imports 80–90% of its food, so 'natural calamity' risk for a bakery is mainly an import-shock risk.")

        scenarios = {
            "No disruption": (0, 0, 0.02),
            "Port congestion": (7, 6, 0.15),
            "Strait chokepoint closure": (18, 14, 0.35),
            "Extreme heat / sandstorm logistics delay": (3, 3, 0.10),
            "Custom": None,
        }
        choice = st.selectbox("Disruption scenario", list(scenarios.keys()), index=2)

        if choice == "Custom":
            colA, colB, colC = st.columns(3)
            delay_days = colA.slider("Lead-time extension (days)", 0, 30, 10)
            cost_premium = colB.slider("Cost premium (%)", 0, 40, 10)
            stockout_prob = colC.slider("Stockout probability", 0.0, 1.0, 0.2, 0.05)
        else:
            delay_days, cost_premium, stockout_prob = scenarios[choice]
            colA, colB, colC = st.columns(3)
            colA.metric("Lead-time extension", f"{delay_days} days")
            colB.metric("Cost premium", f"{cost_premium}%")
            colC.metric("Stockout probability", f"{stockout_prob:.0%}")

        avg_daily_cost = 4200
        buffer_stock_days = max(delay_days, 3)
        buffer_stock_cost = round(buffer_stock_days * avg_daily_cost * 0.02, 0)
        expected_stockout_cost = round(stockout_prob * avg_daily_cost * delay_days * 1.5, 0)

        with st.container(border=True):
            st.markdown(
                f"<div style='padding:6px 4px;'><h4 style='margin-top:0;'>Recommended buffer vs expected stockout cost</h4>"
                f"Holding <b style='color:{COLORS['primary']}'>{buffer_stock_days} days</b> of buffer stock costs approximately "
                f"<b style='color:{COLORS['primary']}'>AED {buffer_stock_cost:,.0f}</b>. The expected cost of not holding buffer, given this scenario's "
                f"probability, is approximately <b style='color:{COLORS['danger']}'>AED {expected_stockout_cost:,.0f}</b>.</div>",
                unsafe_allow_html=True)

            fig2 = go.Figure(go.Bar(
                x=["Cost of holding buffer stock", "Expected cost of stockout"],
                y=[buffer_stock_cost, expected_stockout_cost],
                marker_color=[COLORS["primary"], COLORS["danger"]], marker_line_width=0,
            ))
            fig2.update_layout(**PLOTLY_DARK, height=260, yaxis_title="AED",
                                yaxis=dict(gridcolor=GRID_COLOR), xaxis=dict(gridcolor="rgba(0,0,0,0)"))
            st.plotly_chart(fig2, width='stretch', config={"displayModeBar": False})

    # ---------------- MODULE 3: PANDEMIC PREPAREDNESS ----------------
    with mod3:
        st.markdown("##### 🦠 Pandemic preparedness")
        st.caption("Two linked panels: customer retention/attraction (demand side) and supply chain optimization (supply side).")

        colX, colY = st.columns(2)
        with colX:
            with st.container(border=True):
                st.markdown(f"<h4 style='color:{COLORS['secondary']}'>👥 Customer retention / attraction</h4>", unsafe_allow_html=True)
                repeat_rate = st.slider("Repeat-purchase rate (%)", 0, 100, 62, key="repeat")
                delivery_share = st.slider("Delivery/online-order revenue share (%)", 0, 100, 28, key="delivery")
                basket_size = st.slider("Average basket size (AED)", 20, 200, 74, key="basket")
                flags = []
                if repeat_rate < 50: flags.append("Repeat-purchase rate below healthy baseline")
                if delivery_share < 15: flags.append("Delivery share low — limited resilience to footfall shocks")
                badge, label = ("pill-risk", "Needs attention") if flags else ("pill-ok", "Within healthy range")
                st.markdown(f"<span class='pill {badge}'>{label}</span>", unsafe_allow_html=True)
                for f in flags:
                    st.caption(f"⚠️ {f}")

        with colY:
            with st.container(border=True):
                st.markdown(f"<h4 style='color:{COLORS['primary']}'>🚚 Supply chain optimization</h4>", unsafe_allow_html=True)
                safety_days = st.slider("Safety stock (days of cover)", 0, 30, 9, key="safety")
                alt_suppliers = st.slider("Qualified alternate suppliers per key ingredient", 0, 5, 1, key="alt")
                single_sourced = st.slider("% ingredients single-sourced", 0, 100, 45, key="single")
                flags2 = []
                if safety_days < 7: flags2.append("Safety stock below 7-day resilience threshold")
                if alt_suppliers < 2: flags2.append("Fewer than 2 qualified alternates — concentration risk")
                if single_sourced > 40: flags2.append("Over 40% of ingredients single-sourced")
                badge2, label2 = ("pill-risk", "Needs attention") if flags2 else ("pill-ok", "Within healthy range")
                st.markdown(f"<span class='pill {badge2}'>{label2}</span>", unsafe_allow_html=True)
                for f in flags2:
                    st.caption(f"⚠️ {f}")

        st.caption("Rationale: pandemics create a dual shock — demand shifting from in-store to delivery, and "
                   "supply disruption at the same time — so both sides need to be tracked together as an early-warning system.")

    # ---------------- MODULE 4: SUPPLIER ALTERNATIVES / CONCENTRATION ----------------
    with mod4:
        st.markdown("##### 🌾 Raw material & supplier alternatives")
        st.caption("Herfindahl–Hirschman-style concentration index for core bakery inputs (e.g., flour), "
                   "following the measurement approach used in UAE cereal-supply-risk research (Ali et al., 2022).")

        default_suppliers = pd.DataFrame({
            "Supplier / origin": ["Russia", "Canada", "Australia", "India", "Other"],
            "Spend share (%)": [31, 28, 12, 20, 9],
        })
        st.caption("Edit the spend shares below to reflect current or hypothetical sourcing mix (must sum to ~100%).")
        edited = st.data_editor(default_suppliers, num_rows="fixed", width='stretch', hide_index=True)

        shares = edited["Spend share (%)"].astype(float)
        total = shares.sum()
        shares_norm = shares / total * 100 if total > 0 else shares
        hhi = round((shares_norm ** 2).sum(), 0)

        if hhi < 1500:
            risk_label, risk_badge = "Low concentration risk", "pill-ok"
        elif hhi < 2500:
            risk_label, risk_badge = "Moderate concentration risk", "pill-warn"
        else:
            risk_label, risk_badge = "High concentration risk", "pill-risk"

        colM, colN = st.columns([1, 2])
        with colM:
            with st.container(border=True):
                st.markdown(f"<div class='kpi-label'>Supplier concentration index (HHI)</div>", unsafe_allow_html=True)
                fig_hhi = go.Figure(go.Indicator(
                    mode="gauge+number",
                    value=hhi,
                    number={'font': {'color': COLORS['text'], 'size': 32, 'family': 'Bebas Neue'}},
                    gauge={
                        'axis': {'range': [0, 10000], 'tickcolor': COLORS['text_soft'],
                                 'tickfont': {'color': COLORS['text_soft'], 'size': 9}, 'tickwidth': 1},
                        'bar': {'color': COLORS['primary'], 'thickness': 0.28},
                        'bgcolor': 'rgba(0,0,0,0)',
                        'borderwidth': 0,
                        'steps': [
                            {'range': [0, 1500], 'color': hex_to_rgba(COLORS['success'], 0.20)},
                            {'range': [1500, 2500], 'color': hex_to_rgba(COLORS['warning'], 0.22)},
                            {'range': [2500, 10000], 'color': hex_to_rgba(COLORS['danger'], 0.22)},
                        ],
                        'threshold': {'line': {'color': COLORS['primary_deep'], 'width': 3},
                                      'thickness': 0.85, 'value': hhi},
                    },
                ))
                fig_hhi.update_layout(**PLOTLY_DARK, height=200)
                st.plotly_chart(fig_hhi, width='stretch', config={"displayModeBar": False})
                st.markdown(f"<span class='pill {risk_badge}'>{risk_label}</span>", unsafe_allow_html=True)
                st.caption("Reference bands: <1,500 low · 1,500–2,500 moderate · >2,500 high (standard HHI convention).")

        with colN:
            with st.container(border=True):
                fig3 = go.Figure(go.Pie(labels=edited["Supplier / origin"], values=shares_norm,
                                         hole=0.62, marker=dict(colors=CATEGORICAL, line=dict(color=COLORS["bg"], width=2)),
                                         textfont=dict(color=COLORS["text"], size=11)))
                fig3.update_layout(**PLOTLY_DARK, height=280, legend=dict(font=dict(color=COLORS["text_soft"])),
                                    annotations=[dict(text=f"<b>{hhi:,.0f}</b><br><span style='font-size:10px;color:{COLORS['text_soft']}'>HHI</span>",
                                                       x=0.5, y=0.5, font=dict(size=20, color=COLORS['primary']), showarrow=False)])
                st.plotly_chart(fig3, width='stretch', config={"displayModeBar": False})

        st.markdown("###### Pre-identified alternate suppliers (illustrative)")
        alt_df = pd.DataFrame({
            "Alternate supplier": ["Local UAE mill (blended flour)", "Turkey-origin supplier", "Egypt-origin supplier"],
            "Cost delta vs current": ["+4.5%", "+2.1%", "+1.8%"],
            "Lead-time delta": ["-3 days", "-1 day", "0 days"],
        })
        st.dataframe(alt_df, width='stretch', hide_index=True)

    scenario_context = {
        "inflation": {
            "headline_inf": headline_inf, "food_inf": food_inf, "subsidy_offset": subsidy_offset,
            "base_cost": base_cost, "adj_cost_headline": adj_cost_headline, "adj_cost_food": adj_cost_food,
            "food_cost_pct_adj": food_cost_pct_adj,
        },
        "disruption": {
            "scenario": choice, "delay_days": delay_days, "cost_premium": cost_premium,
            "stockout_prob": stockout_prob, "buffer_stock_cost": buffer_stock_cost,
            "expected_stockout_cost": expected_stockout_cost,
        },
        "pandemic": {
            "repeat_rate": repeat_rate, "delivery_share": delivery_share, "basket_size": basket_size,
            "safety_days": safety_days, "alt_suppliers": alt_suppliers, "single_sourced": single_sourced,
        },
        "supplier": {"hhi": hhi, "risk_label": risk_label, "shares_df": edited, "alt_df": alt_df},
    }
    render_email_share("scenario_resilience", "Email this Scenario & Resilience report",
                        section="scenario_resilience", context=scenario_context)

# ========================================================================
# SECTION C — EMPLOYEE PORTAL
# ========================================================================
elif section == "employee_portal":
    storage_mode = portal_storage_mode()
    badge_cls = "pill-ok" if "Supabase" in storage_mode else "pill-warn"
    st.markdown(f"<span class='pill {badge_cls}'>💾 Storage: {storage_mode}</span>", unsafe_allow_html=True)
    if "Supabase" not in storage_mode:
        with st.expander("⚙️ Set up persistent storage (Supabase)"):
            st.markdown(
                "The portal works right now, but entries only last for this browser session. "
                "To make them persistent across visits and deployments, connect a free Supabase project — "
                "see the **Employee portal setup** section in the README for the exact SQL and secrets to add."
            )

    with st.expander("🏭 Staff breakdown by line", expanded=False):
        dep_df = pd.DataFrame(
            [{"Department": d, "Headcount": c,
              "Focus": DEPARTMENT_PRIORITY_TAG.get(d, ("", ""))[0]} for d, c in DEPARTMENT_HEADCOUNT.items()]
        )
        dep_df.loc[len(dep_df)] = ["Total (itemized bakery lines)", dep_df["Headcount"].sum(), ""]
        st.dataframe(dep_df, width='stretch', hide_index=True)
        cA, cB = st.columns(2)
        cA.metric("Company-wide headcount (current)", COMPANY_HEADCOUNT_CURRENT)
        cB.metric("Growth target (within 1 year)", COMPANY_HEADCOUNT_TARGET,
                   f"+{COMPANY_HEADCOUNT_TARGET - COMPANY_HEADCOUNT_CURRENT} needed")
        st.caption("Company-wide figures per the 27 Jul MoM — larger than the itemized bakery-line total above, "
                   "likely including catering/other roles not broken out per line. Approximate headcounts overall; "
                   "a small illustrative sample of named employees per department is used for the interactive "
                   "demo below — replace with the full roster once available.")

    st.markdown("#### Who are you?")
    col_d1, col_d2 = st.columns([1, 1.4])
    with col_d1:
        dept_choice = st.selectbox("Department", ["🧑‍💼 Manager / HR view"] + list(DEPARTMENT_HEADCOUNT.keys()))

    identity_employee = None
    if dept_choice != "🧑‍💼 Manager / HR view":
        dept_employees = [e for e in EMPLOYEES if e["department"] == dept_choice]
        emp_labels = [f"{e['name']} ({e['id']})" for e in dept_employees]
        with col_d2:
            emp_choice = st.selectbox("Employee", ["— Select —"] + emp_labels)
        if emp_choice != "— Select —":
            identity_employee = dept_employees[emp_labels.index(emp_choice)]

    # -------------------- EMPLOYEE SELF-SERVICE --------------------
    if identity_employee is not None:
        emp_name, emp_id, emp_dept = identity_employee["name"], identity_employee["id"], identity_employee["department"]
        tag_label, tag_cls = DEPARTMENT_PRIORITY_TAG.get(emp_dept, (None, None))
        tag_html = f"<span class='pill {tag_cls}' style='margin-left:6px;'>{tag_label}</span>" if tag_label else ""
        st.markdown(f"<span class='pill pill-ok'>{emp_dept}</span>{tag_html}", unsafe_allow_html=True)

        etab1, etab2, etab3 = st.tabs(["📊 My performance", "📝 Daily log", "🎯 Goals & feedback"])

        # --- My performance ---
        with etab1:
            perf_df, err = portal_fetch("employee_performance", {"employee_id": emp_id})
            if err:
                st.error(f"Could not load performance data: {err}")
            elif perf_df.empty:
                pass
            else:
                for col in ["units_produced", "wastage_pct", "hours_worked", "batch_time_adherence_pct",
                            "quality_pass_pct", "revenue_generated", "break_minutes",
                            "downtime_minutes", "changeover_minutes"]:
                    if col in perf_df.columns:
                        perf_df[col] = pd.to_numeric(perf_df[col], errors="coerce")

                # =========================================================
                # LABOUR EFFICIENCY — §A10 rework. Shows the chain (headline,
                # deduction bar, deduction table, labour economics, benchmark)
                # rather than just a single number, using labour_calc.py so
                # this always matches the Manager/HR view for the same
                # underlying shifts. Employee view leads with performance
                # while working (§A5) — unaffected by equipment failure or
                # changeover, so honest downtime reporting never costs them.
                # =========================================================
                summary = compute_labour_summary(perf_df, emp_dept)
                rolled = summary["rolled"]
                breaks_are_paid = summary["breaks_are_paid"]

                with st.container(border=True):
                    st.markdown(
                        f"<div class='kpi-label'>Performance while working (your number)</div>"
                        f"<div class='kpi-value'>{fmt_pct(rolled['performance_while_working_pct'])}</div>"
                        f"<div style='color:{COLORS['text_soft']}; font-size:0.85rem; margin-top:2px;'>"
                        f"True efficiency (management's number): {fmt_pct(rolled['true_efficiency_pct'])}</div>",
                        unsafe_allow_html=True)
                    st.caption(f"Breaks treated as {'paid' if breaks_are_paid else 'unpaid'} — pending confirmation.")

                    render_deduction_bar_and_table(rolled, breaks_are_paid)
                    render_labour_economics_cards(summary)

                avg_wastage = perf_df["wastage_pct"].mean()
                wastage_pill = ("pill-ok" if avg_wastage <= WASTAGE_TARGET_PCT
                                 else "pill-warn" if avg_wastage <= 3 else "pill-risk")

                c1, c2, c3, c4 = st.columns(4)
                with c1, st.container(border=True):
                    st.markdown(f"<div class='kpi-label'>Avg units / shift</div>"
                                f"<div class='kpi-value' style='font-size:1.7rem;'>{perf_df['units_produced'].mean():.0f}</div>",
                                unsafe_allow_html=True)
                with c2, st.container(border=True):
                    st.markdown(f"<div class='kpi-label'>Avg wastage %</div>"
                                f"<div class='kpi-value' style='font-size:1.7rem;'>{avg_wastage:.1f}%</div>"
                                f"<span class='pill {wastage_pill}'>Target ≤{WASTAGE_TARGET_PCT:.0f}%</span>",
                                unsafe_allow_html=True)
                with c3, st.container(border=True):
                    st.markdown(f"<div class='kpi-label'>Batch-time adherence</div>"
                                f"<div class='kpi-value' style='font-size:1.7rem;'>{perf_df['batch_time_adherence_pct'].mean():.0f}%</div>",
                                unsafe_allow_html=True)
                with c4, st.container(border=True):
                    st.markdown(f"<div class='kpi-label'>Quality pass rate</div>"
                                f"<div class='kpi-value' style='font-size:1.7rem;'>{perf_df['quality_pass_pct'].mean():.0f}%</div>",
                                unsafe_allow_html=True)

                # --- Productivity benchmark (AED 1,000/day per GM's rule of thumb) ---
                c5, c6 = st.columns(2)
                with c5, st.container(border=True):
                    avg_hours = perf_df["hours_worked"].mean() if perf_df["hours_worked"].mean() > 0 else STANDARD_SHIFT_HOURS
                    has_revenue = "revenue_generated" in perf_df.columns and perf_df["revenue_generated"].notna().any()
                    value_per_day = ((perf_df["revenue_generated"].mean() / avg_hours) * STANDARD_SHIFT_HOURS
                                      if has_revenue and avg_hours else None)
                    val_pill = ("pill-ok" if value_per_day is not None and value_per_day >= VALUE_PER_EMPLOYEE_DAY_TARGET
                                else "pill-warn" if value_per_day is not None else "pill-flat")
                    st.markdown(f"<div class='kpi-label'>Value generated / day</div>"
                                f"<div class='kpi-value' style='font-size:1.7rem;'>{fmt_aed(value_per_day)}</div>"
                                f"<span class='pill {val_pill}'>GM benchmark: AED {VALUE_PER_EMPLOYEE_DAY_TARGET:,.0f}/day</span>",
                                unsafe_allow_html=True)
                with c6, st.container(border=True):
                    cap_target = DEPARTMENT_CAPACITY_TARGET.get(emp_dept)
                    if cap_target:
                        util_pct = (perf_df["units_produced"].mean() / cap_target) * 100
                        util_pill = "pill-ok" if util_pct >= 70 else "pill-warn" if util_pct >= 40 else "pill-risk"
                        st.markdown(f"<div class='kpi-label'>Capacity utilization ({emp_dept})</div>"
                                    f"<div class='kpi-value' style='font-size:1.7rem;'>{util_pct:.0f}%</div>"
                                    f"<span class='pill {util_pill}'>vs illustrative target {cap_target}/day</span>",
                                    unsafe_allow_html=True)
                    else:
                        st.markdown(f"<div class='kpi-label'>Capacity utilization</div>"
                                    f"<div class='kpi-value' style='font-size:1rem; color:{COLORS['text_soft']};'>Not tracked for this line</div>",
                                    unsafe_allow_html=True)

                with st.container(border=True):
                    plot_df = perf_df.sort_values("log_date") if "log_date" in perf_df.columns else perf_df
                    if len(plot_df) < 2:
                        st.markdown("<div class='kpi-label'>Units produced trend</div>", unsafe_allow_html=True)
                        st.caption("Log at least 2 shifts to see a trend here.")
                    else:
                        fig_e = go.Figure()
                        fig_e.add_trace(go.Scatter(x=plot_df.get("log_date", plot_df.index), y=plot_df["units_produced"],
                                                    mode="lines+markers", name="Units produced",
                                                    line=dict(color=COLORS["primary"], width=2.5)))
                        fig_e.update_layout(**PLOTLY_DARK, height=260, yaxis=dict(gridcolor=GRID_COLOR),
                                            xaxis=dict(gridcolor="rgba(0,0,0,0)"),
                                            legend=dict(font=dict(color=COLORS["text_soft"])))
                        st.plotly_chart(fig_e, width='stretch', config={"displayModeBar": False})

                # --- Shift breakdown (Arabic Bread runs a separate night shift) ---
                if emp_dept == "Arabic Bread" and "shift" in perf_df.columns and perf_df["shift"].nunique() > 1:
                    with st.container(border=True):
                        st.markdown("###### Morning vs. night shift")
                        shift_summary = perf_df.groupby("shift").agg(
                            avg_units=("units_produced", "mean"),
                            avg_wastage_pct=("wastage_pct", "mean"),
                            shifts_logged=("shift", "count"),
                        ).round(1).reset_index()
                        st.dataframe(shift_summary, width='stretch', hide_index=True)

                st.dataframe(perf_df.drop(columns=[c for c in ["id"] if c in perf_df.columns]),
                             width='stretch', hide_index=True)

        # --- Daily log ---
        with etab2:
            st.caption("Log today's (or a past) shift in plain counts — no percentages to work out yourself. "
                       "Your performance report will still show the percentages, calculated automatically.")
            with st.form(f"daily_log_form_{emp_id}", clear_on_submit=True):
                log_date = st.date_input("Date", value=datetime.now().date())
                if emp_dept == "Arabic Bread":
                    shift = st.selectbox("Shift", ["Morning", "Night"],
                                          help="Arabic Bread runs a separate night shift in addition to the standard morning shift.")
                else:
                    shift = "Morning"
                colf1, colf2 = st.columns(2)
                units = colf1.number_input("Units produced (good units)", 0, 5000, 250, 10)
                units_wasted = colf2.number_input("Units wasted / discarded", 0, 1000, 6, 1,
                                                   help="How many units were thrown away, spoiled, or damaged this shift.")
                hours = colf1.number_input("Hours worked", 0.0, 16.0, STANDARD_SHIFT_HOURS, 0.5)
                units_failed_qc = colf2.number_input("Units that failed QC check", 0, 1000, 5, 1,
                                                       help="How many units failed a quality check this shift.")

                st.markdown("###### Break, downtime & changeover")
                st.caption(
                    "These are used to work out true efficiency. Reporting downtime honestly "
                    "does not lower your personal performance score — it flags an equipment "
                    "or scheduling issue."
                )
                colf5, colf6, colf7 = st.columns(3)
                break_minutes = colf5.number_input(
                    "Break time taken (minutes)", 0, 600, LABOUR_CONFIG["default_break_minutes"], 5,
                    help="Total break minutes this shift.")
                downtime_minutes = colf6.number_input(
                    "Waiting for equipment (minutes)", 0, 600, 0, 5,
                    help="Machine down, oven not ready, or waiting on materials.")
                changeover_minutes = colf7.number_input(
                    "Changeover / setup (minutes)", 0, 600, 0, 5,
                    help="Switching between products or batch setup.")

                colf3, colf4 = st.columns(2)
                batches_completed = colf3.number_input("Batches completed", 0, 200, 10, 1)
                batches_on_time = colf4.number_input("Batches finished on/within standard time", 0, 200, 9, 1,
                                                       help="Of the batches completed, how many finished within the standard time for that batch?")
                revenue = st.number_input(
                    "Revenue/value generated (AED, optional)", min_value=0.0, max_value=50000.0, value=None, step=50.0,
                    help="Feeds the productivity benchmark: GM's rule of thumb is each employee should "
                         "generate ~AED 1,000/day."
                )
                st.caption("Leave blank if not tracked for your line — it will be excluded rather than counted as zero.")
                notes = st.text_area("Notes (optional)", placeholder="Any incidents, equipment issues, etc.")
                submitted = st.form_submit_button("Submit shift log", width='stretch')
                if submitted:
                    total_output = units + units_wasted
                    wastage_pct = round((units_wasted / total_output) * 100, 1) if total_output > 0 else 0.0
                    batch_adh_pct = round((batches_on_time / batches_completed) * 100) if batches_completed > 0 else 0
                    quality_pct = round(max(units - units_failed_qc, 0) / units * 100) if units > 0 else 0
                    ok, err = portal_insert("employee_performance", {
                        "employee_name": emp_name, "employee_id": emp_id, "department": emp_dept,
                        "log_date": str(log_date), "shift": shift, "units_produced": int(units),
                        "units_wasted": int(units_wasted), "batches_completed": int(batches_completed),
                        "batches_on_time": int(batches_on_time), "units_failed_qc": int(units_failed_qc),
                        "wastage_pct": wastage_pct, "hours_worked": float(hours),
                        "batch_time_adherence_pct": batch_adh_pct, "quality_pass_pct": quality_pct,
                        "break_minutes": int(break_minutes), "downtime_minutes": int(downtime_minutes),
                        "changeover_minutes": int(changeover_minutes),
                        "revenue_generated": (float(revenue) if revenue is not None else None), "notes": notes,
                    })
                    if ok:
                        st.toast(f"Shift log saved — wastage {wastage_pct}%, batch adherence {batch_adh_pct}%, "
                                 f"quality pass {quality_pct}% (calculated automatically from your counts).", icon="✅")
                        st.rerun()
                    else:
                        st.error(f"Could not save: {err}")

        # --- Goals & feedback ---
        with etab3:
            goals_df, err = portal_fetch("employee_goals", {"employee_id": emp_id})
            if err:
                st.error(f"Could not load goals: {err}")
            elif not goals_df.empty:
                for _, row in goals_df.iterrows():
                    with st.container(border=True):
                        st.markdown(f"**Goal:** {row.get('goal_text', '—')}")
                        st.caption(f"Self-assessment: {row.get('self_assessment') or '—'}")
                        if row.get("manager_feedback"):
                            st.markdown(f"<span class='pill pill-ok'>Manager feedback</span> {row['manager_feedback']}",
                                        unsafe_allow_html=True)
            else:
                pass

            st.markdown("###### Add a new goal / self-assessment")
            with st.form(f"goals_form_{emp_id}", clear_on_submit=True):
                goal_text = st.text_area("Current goal / focus area")
                self_assessment = st.text_area("Self-assessment for this period")
                submitted_g = st.form_submit_button("Save", width='stretch')
                if submitted_g:
                    ok, err = portal_insert("employee_goals", {
                        "employee_name": emp_name, "employee_id": emp_id, "department": emp_dept,
                        "goal_text": goal_text, "self_assessment": self_assessment,
                        "manager_feedback": None,
                    })
                    if ok:
                        st.toast("Goal saved.", icon="✅")
                        st.rerun()
                    else:
                        st.error(f"Could not save: {err}")

    # -------------------- MANAGER / HR VIEW --------------------
    elif dept_choice == "🧑‍💼 Manager / HR view":
        all_perf, err = portal_fetch("employee_performance")
        if err:
            st.error(f"Could not load team data: {err}")
        elif all_perf.empty:
            pass
        else:
            for col in ["units_produced", "wastage_pct", "batch_time_adherence_pct", "quality_pass_pct", "revenue_generated", "hours_worked"]:
                if col in all_perf.columns:
                    all_perf[col] = pd.to_numeric(all_perf[col], errors="coerce")

            st.markdown("#### Performance comparison within a department")
            available_depts = sorted(all_perf["department"].dropna().unique()) if "department" in all_perf.columns else []
            if not available_depts:
                pass
            else:
                dept_filter = st.selectbox("Choose a department to compare", available_depts, key="mgr_dept_filter")
                tag_label, tag_cls = DEPARTMENT_PRIORITY_TAG.get(dept_filter, (None, None))
                tag_html = f"<span class='pill {tag_cls}' style='margin-left:8px;'>{tag_label}</span>" if tag_label else ""

                dept_perf = all_perf[all_perf["department"] == dept_filter].copy()

                # =====================================================
                # LABOUR EFFICIENCY — department rollup (§A5, §A8). Manager
                # view leads with true efficiency (the cost view — downtime
                # correctly reduces it, since Grandiose paid for that time).
                # Hours are aggregated FIRST, then percentages are computed
                # from the aggregated hours — averaging each employee's own
                # percentage is wrong whenever employees logged different
                # numbers of shifts, and it is a silent error.
                # =====================================================
                dept_summary_labour = compute_labour_summary(dept_perf, dept_filter)
                dept_rolled = dept_summary_labour["rolled"]

                with st.container(border=True):
                    st.markdown(
                        f"<div class='kpi-label'>True efficiency — {dept_filter} (management's number)</div>"
                        f"<div class='kpi-value'>{fmt_pct(dept_rolled['true_efficiency_pct'])}</div>"
                        f"<div style='color:{COLORS['text_soft']}; font-size:0.85rem; margin-top:2px;'>"
                        f"Naive mean of each employee's own %: {fmt_pct(dept_summary_labour['naive_mean_true_efficiency'])} — "
                        f"wrong whenever employees logged different numbers of shifts</div>",
                        unsafe_allow_html=True)
                    st.caption(f"Breaks treated as {'paid' if dept_summary_labour['breaks_are_paid'] else 'unpaid'} "
                               f"— pending confirmation.")
                    render_labour_economics_cards(dept_summary_labour)

                def _value_per_day(row_group):
                    hrs = row_group["hours_worked"].mean()
                    hrs = hrs if hrs and hrs > 0 else STANDARD_SHIFT_HOURS
                    rev = row_group["revenue_generated"].mean() if "revenue_generated" in row_group.columns else 0.0
                    return (rev / hrs) * STANDARD_SHIFT_HOURS

                dept_summary = dept_perf.groupby("employee_name").agg(
                    shifts_logged=("employee_name", "count"),
                    avg_units=("units_produced", "mean"),
                    avg_wastage_pct=("wastage_pct", "mean"),
                    avg_batch_adherence=("batch_time_adherence_pct", "mean"),
                    avg_quality_pass=("quality_pass_pct", "mean"),
                ).round(1).reset_index()
                dept_summary["value_per_day_aed"] = dept_summary["employee_name"].apply(
                    lambda n: round(_value_per_day(dept_perf[dept_perf["employee_name"] == n]), 0))
                dept_summary["wastage_status"] = dept_summary["avg_wastage_pct"].apply(
                    lambda w: "✅ On target" if w <= WASTAGE_TARGET_PCT else "⚠️ Above 1% target")

                # Composite performance score (0-100) so "who's doing best" has a
                # single clear answer instead of reading several charts separately.
                max_units = dept_summary["avg_units"].max()
                dept_summary["units_normalized"] = (dept_summary["avg_units"] / max_units * 100) if max_units > 0 else 0
                dept_summary["performance_score"] = (
                    dept_summary["avg_quality_pass"] * 0.30
                    + dept_summary["avg_batch_adherence"] * 0.30
                    + (100 - dept_summary["avg_wastage_pct"]).clip(lower=0) * 0.20
                    + dept_summary["units_normalized"] * 0.20
                ).round(1)
                dept_summary = dept_summary.sort_values("performance_score", ascending=False).reset_index(drop=True)
                dept_summary["rank"] = dept_summary.index + 1

                cap_target = DEPARTMENT_CAPACITY_TARGET.get(dept_filter)

                with st.container(border=True):
                    st.markdown(f"**{dept_filter}**{tag_html} — {DEPARTMENT_HEADCOUNT.get(dept_filter, '—')} total headcount "
                                f"on the line, {dept_perf['employee_name'].nunique()} with logged shifts",
                                unsafe_allow_html=True)
                    if cap_target:
                        util_pct = (dept_summary["avg_units"].mean() / cap_target) * 100
                        st.caption(f"Capacity utilization vs illustrative target of {cap_target} units/day/employee: **{util_pct:.0f}%**")
                    if len(dept_summary) > 0:
                        top = dept_summary.iloc[0]
                        st.markdown(f"<span class='pill pill-ok'>🏆 Top performer: {top['employee_name']} "
                                    f"(score {top['performance_score']:.0f}/100)</span>", unsafe_allow_html=True)
                    display_cols = ["rank", "employee_name", "performance_score", "shifts_logged", "avg_units",
                                     "avg_wastage_pct", "avg_batch_adherence", "avg_quality_pass",
                                     "value_per_day_aed", "wastage_status"]
                    st.dataframe(dept_summary[display_cols], width='stretch', hide_index=True)
                    st.caption("Performance score = 30% quality pass rate + 30% batch-time adherence + "
                               "20% (100 − wastage %) + 20% units produced (normalized to the department's top performer).")
                    render_email_share(
                        f"dept_{dept_filter}", f"Email the {dept_filter} comparison",
                        section="employee_department",
                        context={
                            "department": dept_filter,
                            "headcount": DEPARTMENT_HEADCOUNT.get(dept_filter, "—"),
                            "tag_label": tag_label,
                            "summary_df": dept_summary[display_cols],
                        },
                    )

                with st.container(border=True):
                    st.markdown("###### Comparative progress across employees in this department")
                    fig_cmp = go.Figure()
                    fig_cmp.add_trace(go.Bar(x=dept_summary["employee_name"], y=dept_summary["avg_quality_pass"],
                                              name="Quality pass %", marker_color=COLORS["primary"], marker_line_width=0))
                    fig_cmp.add_trace(go.Bar(x=dept_summary["employee_name"], y=dept_summary["avg_batch_adherence"],
                                              name="Batch-time adherence %", marker_color=COLORS["tertiary"], marker_line_width=0))
                    fig_cmp.add_trace(go.Bar(x=dept_summary["employee_name"],
                                              y=(100 - dept_summary["avg_wastage_pct"]).clip(lower=0),
                                              name="100 − wastage %", marker_color=COLORS["secondary"], marker_line_width=0))
                    fig_cmp.update_layout(**PLOTLY_DARK, height=300, barmode="group", yaxis_title="%",
                                           yaxis=dict(gridcolor=GRID_COLOR, range=[0, 105]),
                                           xaxis=dict(gridcolor="rgba(0,0,0,0)"),
                                           legend=dict(orientation="h", y=-0.15, font=dict(color=COLORS["text_soft"])))
                    st.plotly_chart(fig_cmp, width='stretch', config={"displayModeBar": False})

                with st.container(border=True):
                    fig_dept = go.Figure()
                    fig_dept.add_trace(go.Bar(x=dept_summary["employee_name"], y=dept_summary["avg_units"],
                                               name="Avg units/shift", marker_color=COLORS["primary"], marker_line_width=0))
                    fig_dept.update_layout(**PLOTLY_DARK, height=260, yaxis_title="Avg units / shift",
                                            yaxis=dict(gridcolor=GRID_COLOR), xaxis=dict(gridcolor="rgba(0,0,0,0)"))
                    st.plotly_chart(fig_dept, width='stretch', config={"displayModeBar": False})

                with st.container(border=True):
                    fig_dept2 = go.Figure(go.Bar(x=dept_summary["employee_name"], y=dept_summary["avg_wastage_pct"],
                                                  marker_color=COLORS["secondary"], marker_line_width=0))
                    fig_dept2.add_hline(y=WASTAGE_TARGET_PCT, line_dash="dot", line_color=COLORS["text_soft"],
                                         annotation_text="1% target", annotation_font_color=COLORS["text_soft"])
                    fig_dept2.update_layout(**PLOTLY_DARK, height=260, yaxis_title="Avg wastage %",
                                             yaxis=dict(gridcolor=GRID_COLOR), xaxis=dict(gridcolor="rgba(0,0,0,0)"))
                    st.plotly_chart(fig_dept2, width='stretch', config={"displayModeBar": False})

                if dept_filter == "Arabic Bread" and "shift" in dept_perf.columns and dept_perf["shift"].nunique() > 1:
                    with st.container(border=True):
                        st.markdown("###### Morning vs. night shift — Arabic Bread")
                        shift_cmp = dept_perf.groupby("shift").agg(
                            avg_units=("units_produced", "mean"),
                            avg_wastage_pct=("wastage_pct", "mean"),
                            shifts_logged=("shift", "count"),
                        ).round(1).reset_index()
                        st.dataframe(shift_cmp, width='stretch', hide_index=True)

            st.markdown("#### All departments — overview")
            dept_overview = all_perf.groupby("department").agg(
                shifts_logged=("department", "count"),
                avg_units=("units_produced", "mean"),
                avg_wastage_pct=("wastage_pct", "mean"),
            ).round(1).reset_index() if "department" in all_perf.columns else pd.DataFrame()
            if not dept_overview.empty:
                dept_overview["headcount"] = dept_overview["department"].map(DEPARTMENT_HEADCOUNT)
                dept_overview["capacity_target"] = dept_overview["department"].map(DEPARTMENT_CAPACITY_TARGET)
                dept_overview["utilization_pct"] = (dept_overview["avg_units"] / dept_overview["capacity_target"] * 100).round(0)
                dept_overview["focus"] = dept_overview["department"].map(lambda d: DEPARTMENT_PRIORITY_TAG.get(d, ("", ""))[0])
                with st.container(border=True):
                    st.dataframe(dept_overview.sort_values("avg_units", ascending=False), width='stretch', hide_index=True)

        st.markdown("#### Recent goals & self-assessments")
        all_goals, err = portal_fetch("employee_goals")
        if err:
            st.error(f"Could not load goals: {err}")
        elif all_goals.empty:
            pass
        else:
            for _, row in all_goals.iterrows():
                with st.container(border=True):
                    st.markdown(f"**{row.get('employee_name','—')}** ({row.get('department','—')}) — {row.get('goal_text','—')}")
                    st.caption(f"Self-assessment: {row.get('self_assessment') or '—'}")
                    if row.get("manager_feedback"):
                        st.markdown(f"<span class='pill pill-ok'>Feedback given</span> {row['manager_feedback']}",
                                    unsafe_allow_html=True)

        st.markdown("###### Give feedback on an employee's goal")
        with st.form("manager_feedback_form", clear_on_submit=True):
            fb_labels = [f"{e['department']} — {e['name']}" for e in EMPLOYEES]
            fb_choice = st.selectbox("Employee", fb_labels, key="fb_employee")
            fb_text = st.text_area("Feedback")
            fb_submit = st.form_submit_button("Submit feedback", width='stretch')
            if fb_submit:
                fb_emp = EMPLOYEES[fb_labels.index(fb_choice)]
                ok, err = portal_insert("employee_goals", {
                    "employee_name": fb_emp["name"], "employee_id": fb_emp["id"], "department": fb_emp["department"],
                    "goal_text": "(Manager feedback entry)", "self_assessment": None,
                    "manager_feedback": fb_text,
                })
                if ok:
                    st.toast("Feedback recorded.", icon="✅")
                    st.rerun()
                else:
                    st.error(f"Could not save: {err}")
    else:
        st.info("Select your department and name (or Manager / HR view) above to get started.")

# ========================================================================
# SECTION D — COMPANY PROFILE
# ========================================================================
elif section == "company_profile":
    COMPANY_INFO = [
        ("🏭", "Production capacity & growth", [
            "Flour Country Bakery is currently running at only 30-35% utilization — roughly 65% headroom still available.",
            "Target is 100% utilization within a year, with an estimated AED 10M/month production potential once reached.",
            "Granola Catering is running at approximately 50% utilization.",
            "Company-wide headcount is ~110 today, with a growth target of 200 within a year.",
        ]),
        ("⏰", "Shift structure", [
            "Standard shift is 9 hours, including a break.",
            "Arabic Bread is the one line that also runs a separate night shift, in addition to the standard morning shift.",
        ]),
        ("🛒", "Sales channel mix (B2B vs. B2C)", [
            "Flour Country Bakery is 100% B2B.",
            "Catering is roughly 35% B2C and 65% B2B.",
        ]),
        ("💰", "Cost structure & targets", [
            "Food cost: capped at a maximum of 30%.",
            "Labour cost: currently running at ~40% — target is 16-18%, with 20% as the absolute ceiling.",
            "Utilities: capped at a maximum of 10%.",
            "Rent: capped at a maximum of 15%.",
            "Net profit target: 20-25%.",
            "Total expense ceiling (all costs combined): 75-80% of revenue.",
        ]),
        ("🏷️", "Pricing approach", [
            "Material cost is scaled 35-50% depending on client volume.",
            "Supermarket supply is priced cost-to-cost, with no margin built in.",
            "B2B contract pricing varies by negotiation and account.",
        ]),
        ("📈", "Productivity benchmark", [
            "No formal labour productivity measurement system exists yet.",
            "The GM's working rule of thumb: each employee should generate approximately AED 1,000/day of value.",
            "This is the benchmark now built into the Employee Portal's productivity tracking.",
        ]),
        ("🔥", "Top & underutilized product lines", [
            "Top revenue/profit drivers: daily bread (Arabic bread, toasted bread, buns, Samurai sandwich) and pastry.",
            "Least-utilized lines: Tahina and Turkish baklava — both flagged as ramp-up focus areas.",
        ]),
        ("📦", "Inventory management", [
            "Stock is managed on a FIFO (first-in, first-out) basis.",
            "Ordering cycles: bi-monthly for shelf-stable/dry goods, weekly for short-shelf-life items, daily for fresh fruit and vegetables.",
            "Payment terms: ~95% of suppliers are on net 90 days (60 days from statement of account); a few critical items run 30-60 days.",
            "Approved vendor base: ~65-70 for bakery, ~160-170 for catering.",
            "Safety stock normally aligns to supplier lead times, extending to ~3 months during geopolitical disruptions.",
            "Packaging (sourced from China/Indonesia): ~1 month of stock held against a 15-day lead time.",
            "De-risking in progress: alternate vendors are being onboarded for critical ingredients.",
            "KPIs monitored: inventory movement (fast/slow movers), depletion cycles, and overstock/obsolescence — reviewed bi-weekly.",
        ]),
        ("🗑️", "Wastage management", [
            "Current wastage: 2-3% — target is 1%.",
            "Plan to close the gap: grow sales volume and push supermarket promotions on affected items.",
            "Near-expiry raw materials are redirected into production rather than discarded.",
            "Near-expiry finished goods (with shelf life over 30 days) move to clearance pricing.",
            "Fully expired stock is written off.",
            "Yield: catering runs ~90-95% (about 5% loss to spillage/damage); bakery yield loss is minimal; frozen raw meat yield loss is 20-25%.",
        ]),
        ("🗓️", "Dashboard & next steps (from this meeting)", [
            "ERP integration is planned so the dashboard can eventually auto-update from live company data.",
            "In-dashboard email sharing was requested — ✅ already built into this dashboard.",
            "An employee productivity module linking individual output to section-level sales was requested — ✅ already built into the Employee Portal.",
            "A separate, equivalent dashboard is planned for Granola Catering.",
            "The GM wants the dashboard link shared for internal review before Friday.",
            "Next steps: send Friday feedback slots, share the dashboard link, and continue refining with confirmed figures ahead of the mid-review.",
        ]),
    ]

    for icon, title, points in COMPANY_INFO:
        with st.container(border=True):
            rows = "".join([f"<div class='info-row'><span>{p}</span></div>" for p in points])
            st.markdown(f"<div style='padding:14px 18px;'><h4 style='margin:0 0 10px 0; font-size:1.05rem;'>"
                        f"{icon} {title}</h4>{rows}</div>", unsafe_allow_html=True)

    st.caption("Source: personal meeting with the Grandiose GM, 27 Jul. Figures here are as reported in that "
               "conversation and should be treated as the current source of truth pending any further updates.")
    render_email_share("company_profile", "Email the Company Profile", section="company_profile",
                        context={"company_info": COMPANY_INFO})

# ========================================================================
# SECTION E — DATA PROCESSOR (AI-powered rough-data cleanup)
# ========================================================================
# ========================================================================
# SECTION E — BAKERY PRODUCT PERFORMANCE (SKU)
# ========================================================================
elif section == "sku_performance":
    products = load_sku_products()
    k = sku_data.kpis(products)

    # --- KPI strip ---
    with st.container(border=True):
        kpi_cols = st.columns(5, vertical_alignment="top")
        # Cards come from sku_data as label / value / sub-line / trend, so
        # the page only decides sizing and colour. Figures share one size;
        # the two name KPIs step down so a long product name wraps to two
        # lines without pushing the row around, and the cell has a fixed
        # min-height so all five labels stay on one line.
        for col, card in zip(kpi_cols, sku_data.kpi_cards(products)):
            size = "1.4rem" if card["trend"] is None else "2.1rem"
            colour = COLORS["primary"] if card.get("accent") else COLORS["text"]
            sub_colour = {
                "up": COLORS["primary"],
                "down": COLORS["danger"],
            }.get(card["trend"], COLORS["text_soft"])
            arrow = {"up": "↑ ", "down": "↓ "}.get(card["trend"], "")
            col.markdown(
                f"<div style='padding:14px 14px 10px 14px; min-height:126px;'>"
                f"<div class='kpi-label'>{card['label']}</div>"
                f"<div class='kpi-value' style='font-size:{size}; color:{colour}; "
                f"line-height:1.15;'>{card['value']}</div>"
                f"<div style='color:{sub_colour}; font-size:0.72rem; font-weight:600; "
                f"margin-top:6px;'>{arrow}{card['sub']}</div>"
                f"</div>",
                unsafe_allow_html=True,
            )

    st.markdown("#### Division performance")
    st.caption("Each bubble is one SKU, sized by units sold and grouped into its division. "
               "Hover for detail, or click a bubble to open that SKU in the table below.")

    division_choice = st.radio(
        "Division filter",
        ["All", *sku_data.DIVISIONS],
        horizontal=True, label_visibility="collapsed", key="sku_division_filter",
    )
    visible = products if division_choice == "All" else \
        [p for p in products if p["division"] == division_choice]

    selected_sku = st.session_state.get("sku_selected")
    # A SKU filtered out of view should not stay selected.
    if selected_sku and not any(p["sku"] == selected_sku for p in visible):
        selected_sku = None
        st.session_state["sku_selected"] = None

    with st.container(border=True):
        event = st.plotly_chart(
            build_bubble_figure(visible, selected_sku),
            width='stretch', key="sku_bubbles",
            on_select="rerun", selection_mode="points",
            config={"displayModeBar": False},
        )

    # Clicking a bubble selects that SKU; the table below reacts to it.
    picked = (event.selection or {}).get("points") if event else None
    if picked:
        clicked = picked[0].get("customdata", [None] * 7)[6]
        if clicked and clicked != selected_sku:
            st.session_state["sku_selected"] = clicked
            st.rerun()

    if selected_sku:
        chosen = next(p for p in visible if p["sku"] == selected_sku)
        pick_cols = st.columns([6, 1], vertical_alignment="center")
        pick_cols[0].markdown(
            f"<span class='pill pill-ok'>Selected · {chosen['product']} "
            f"({chosen['sku']}) — {chosen['division']}</span>",
            unsafe_allow_html=True,
        )
        if pick_cols[1].button("Clear", key="sku_clear", width='stretch'):
            st.session_state["sku_selected"] = None
            st.rerun()

    # --- SKU details ---
    st.markdown("#### SKU details")
    query = st.text_input(
        "Search SKU or product", placeholder="Search SKU or product…",
        label_visibility="collapsed", key="sku_search",
    ).strip().lower()

    matching = [p for p in visible
                if not query
                or query in p["sku"].lower() or query in p["product"].lower()]
    if query and not matching:
        st.caption(f"No SKU or product matches “{query}”.")

    selected_division = next(
        (p["division"] for p in visible if p["sku"] == selected_sku), None)

    for division in sku_data.DIVISIONS:
        rows = [p for p in matching if p["division"] == division]
        if not rows:
            continue
        title = division
        if division == sku_data.SEASONAL_DIVISION:
            title = f"{division} — {sku_data.SEASONAL_COLLECTION}"
        # Open the division holding the clicked SKU, or every division with
        # a search hit; otherwise leave them closed.
        expanded = (division == selected_division) or bool(query)
        with st.expander(f"{title}  ·  {len(rows)} SKUs", expanded=expanded):
            frame = sku_table(rows, division)
            if selected_sku and selected_sku in set(frame["SKU"]):
                highlight = hex_to_rgba(COLORS["primary"], 0.16)
                styled = frame.style.apply(
                    lambda row: [f"background-color: {highlight}"
                                 if row["SKU"] == selected_sku else "" for _ in row],
                    axis=1,
                )
                st.dataframe(styled, width='stretch', hide_index=True)
            else:
                st.dataframe(frame, width='stretch', hide_index=True)

    st.caption("Figures are illustrative pending Grandiose-provided SKU actuals. "
               "Contribution % is each SKU's share of total bakery sales; rank is by "
               "sales value across all divisions.")

elif section == "data_processor":
    ai_client = get_anthropic_client()
    badge_cls = "pill-ok" if ai_client is not None else "pill-warn"
    ai_status = "AI processing: connected" if ai_client is not None else "AI processing: not configured"
    st.markdown(f"<span class='pill {badge_cls}'>🤖 {ai_status}</span>", unsafe_allow_html=True)
    if ai_client is None:
        with st.expander("⚙️ Set up AI processing (Anthropic API key)"):
            st.markdown(
                "This feature reads whatever's in your uploaded files — even messy or inconsistently "
                "formatted data — and turns it into a cleaned, evaluated Excel report. It needs an "
                "Anthropic API key. See the **Data processor setup** section in the README for where "
                "to get one and exactly what to add to Streamlit Cloud's Secrets panel."
            )

    st.markdown("#### Upload files")
    uploaded_files = st.file_uploader(
        "Drop Excel, PDF, or Word files here",
        type=["xlsx", "xls", "csv", "pdf", "docx"],
        accept_multiple_files=True,
        label_visibility="collapsed",
    )

    if uploaded_files:
        st.caption(f"{len(uploaded_files)} file(s) ready: " + ", ".join(f.name for f in uploaded_files))

    process_clicked = st.button("Process files", width='stretch', disabled=not uploaded_files, type="primary")

    if process_clicked and uploaded_files:
        with st.spinner("Reading files and asking the AI to interpret and evaluate the data..."):
            file_contents = {}
            for f in uploaded_files:
                fbytes = f.getvalue()
                ext = f.name.lower().rsplit(".", 1)[-1]
                if ext == "pdf":
                    file_contents[f.name] = extract_pdf_text(fbytes)
                elif ext == "docx":
                    file_contents[f.name] = extract_docx_text(fbytes)
                else:
                    file_contents[f.name] = extract_excel_text(fbytes, f.name)
            result, err = interpret_uploaded_files(file_contents)
        if err:
            st.error(err)
            st.session_state.pop("smart_upload_result", None)
        else:
            st.session_state["smart_upload_result"] = result
            st.session_state["smart_upload_files"] = [f.name for f in uploaded_files]
            st.toast("Files processed.", icon="✅")

    result = st.session_state.get("smart_upload_result")
    if result:
        st.markdown("#### Results")
        with st.container(border=True):
            st.markdown(f"**Detected data type:** {result.get('detected_data_type', '—')}")
            st.caption(result.get("summary", ""))

        for sheet_info in result.get("sheets", []):
            with st.container(border=True):
                st.markdown(f"###### {sheet_info.get('title', sheet_info.get('sheet_name', 'Sheet'))}")
                columns = sheet_info.get("columns", [])
                rows = sheet_info.get("rows", [])
                if columns and rows:
                    preview_df = pd.DataFrame(
                        [r + [""] * (len(columns) - len(r)) if len(r) < len(columns) else r[:len(columns)] for r in rows],
                        columns=columns,
                    )
                    st.dataframe(preview_df, width='stretch', hide_index=True)
                for ins in sheet_info.get("insights", []):
                    st.caption(f"💡 {ins}")

        excel_bytes = build_excel_report("smart_upload", context={
            "result": result, "source_files": st.session_state.get("smart_upload_files", []),
        })
        st.download_button(
            "Download processed Excel report", data=excel_bytes,
            file_name=f"Grandiose_Bakery_Processed_Data_{datetime.now().strftime('%Y%m%d')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            width='stretch', type="primary",
        )
        render_email_share("smart_upload", "Email this processed report", section="smart_upload",
                            context={"result": result, "source_files": st.session_state.get("smart_upload_files", [])})

    st.caption("Files are sent to Anthropic's API for interpretation and are not stored by this dashboard "
               "beyond the current session. Numbers are AI-interpreted from whatever was uploaded — "
               "always sanity-check before acting on them.")

render_footer()
